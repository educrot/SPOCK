#code for long term scheduling for SPECULOOS
from alive_progress import alive_bar
from astropy.table import Table
from astropy import units as u
from astropy.coordinates import SkyCoord, get_sun, AltAz, EarthLocation
from astropy.utils import iers
from astropy.time import Time
from astropy.utils.data import clear_download_cache
clear_download_cache()
from astroplan import FixedTarget, AltitudeConstraint, MoonSeparationConstraint,AtNightConstraint,observability_table, is_observable, months_observable,time_grid_from_range,LocalTimeConstraint, is_always_observable
from astroplan import TimeConstraint,Observer,moon_illumination
from colorama import Fore
from datetime import date , datetime , timedelta
from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from eScheduler.spe_schedule import SPECULOOSScheduler, Schedule, ObservingBlock, Transitioner
import functools
from functools import reduce
#import gspread
import numpy as np
#from oauth2client.service_account import ServiceAccountCredentials
import os
import os.path, time
import pkg_resources
import pandas as pd
import paramiko
import requests
import shutil
import ssl
import subprocess
import sys
from tqdm.auto import tqdm
import yaml
from SPOCK import pwd_appcs,pwd_HUB,user_portal,pwd_portal,pwd_appcs,pwd_SNO_Reduc1,user_chart_studio,pwd_chart_studio,path_spock

iers.IERS_A_URL  = 'https://datacenter.iers.org/data/9/finals2000A.all' #'http://maia.usno.navy.mil/ser7/finals2000A.all'#'ftp://cddis.gsfc.nasa.gov/pub/products/iers/finals2000A.all'
#from astroplan import download_IERS_A
#download_IERS_A()
ssl._create_default_https_context = ssl._create_unverified_context


import SPOCK.ETC as ETC
from SPOCK.make_night_plans import make_np
import SPOCK.upload_night_plans as SPOCKunp

def get_hours_files_SNO(username = 'speculoos',password = pwd_SNO_Reduc1):
    """ get nb hours obs on SNO

    Returns
    -------
    txt file
        file ObservationHours.txt

    """
    hostname = '172.16.3.11'
    port = 22
    #if __name__ == "__main__":
    paramiko.util.log_to_file('paramiko.log')
    s = paramiko.SSHClient()
    s.load_system_host_keys()
    try:
        s.connect(hostname, port, username, password)
    except TimeoutError:
        sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' Make sure the VPN is connected')

    ftp_client=s.open_sftp()
    ftp_client.get('/home/speculoos/SNO/ObservationHours/ObservationHours.txt',path_spock + '/survey_hours/ObservationHours.txt')
    ftp_client.close()

    s.close()
    df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',',skipinitialspace=True)
    df = df.sort_values(['Target'])
    df.to_csv(path_spock + '/survey_hours/ObservationHours.txt',sep=',',index=False)

def max_unit_list(x):
    """ return max of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        max of list

    """
    a = max(x)
    return a.value

def first_elem_list(x):
    """ first element of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        first element of list

    """
    a = x[0]
    return a

def last_elem_list(x):
    """ last element of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        last element

    """
    a = x[-1]
    return a

def coord_transfotm_to_alt(x,frame):
    """ transform astropy.coordinates to a given frame

    Parameters
    ----------
    x  : astropy.coordinates
    coordinates
    frame : str
        frame in astropy.coordinates

    Returns
    -------
    astropy.coordinates
        coordinates in the frame chosen

    """
    a = (x.coord.transform_to(frame).alt)
    return a

def index_list1_list2(list1, list2): #list 2 longer than list 1
    """ index of list1 in list2 and list2 in list1

    Parameters
    ----------
    list1 : list

    list2 : list

    Returns
    -------
    list
        list of index of list1 in list2 and list2 in list1

    """
    idx_list1_in_list2 = []
    idx_list2_in_list1 = []
    for i in range(len(list2)):
        for j in range(len(list1)):
            if list2[i] == list1[j]:
                idx_list1_in_list2.append(i)
                idx_list2_in_list1.append(j)
    return idx_list1_in_list2, idx_list2_in_list1

def Diff_list(li1, li2):
    """ Inform on the difference between two lists

    Parameters
    ----------
    li1: list
    li2: list

    Returns
    -------
    list
        Elements than are in list 1 but not in list 2
    """

    return (list(set(li1) - set(li2)))

def compare_target_lists(path_target_listuser, user = 'educrot',password = pwd_portal):
    """ Compare the target list from the given folder to the one on STARGATE and Cambridge server
    If different trigger a warning a tell how many targets are actually different from the referenced target list

    Parameters
    ----------
    path_target_list: str
        path on your computer toward the target list, by default take the one on the Cambridge server

    Returns
    -------
    print
         An idication about if the target list is the referenced one or not

    """
    TargetURL="http://www.mrao.cam.ac.uk/SPECULOOS/target_list_gaia.csv"
    resp = requests.get(TargetURL, auth=(user, password))
    open('target_list_gaia.csv', 'wb').write(resp.content)
    content = resp.text.replace("\n", "")
    df_target_list_gaia = pd.read_csv('target_list_gaia.csv', delimiter=',',skipinitialspace=True,error_bad_lines=False)
    df_user = pd.read_csv(path_target_list, delimiter=' ',skipinitialspace=True,error_bad_lines=False)
    if Diff_list(df_user['Name'],df_target_list_gaia['spc']):
        print(Fore.YELLOW + 'WARNING:  ' + Fore.BLACK + ' Tragets in User\'s list but not in Cambridge server\'s list: ',Diff_list(df_user['Name'],df_target_list_gaia['spc'] ))
    if Diff_list(df_target_list_gaia['spc'],df_user['Name'] ):
        print(Fore.YELLOW + 'WARNING:  ' + Fore.BLACK + ' Tragets in Cambridge server\'s list but not in User\'s list: ',Diff_list(df_target_list_gaia['spc'],df_user['Name'] ))
    else:
        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' OK ! User\'s list is similar to the one on the Cambridge server')

def SSO_planned_targets(date,telescope):
    """ tell which target are scheduled on SSO on a given day

    Parameters
    ----------
    date : date
        date of day in fmt 'yyyy-mm-dd'
    telescope : str
        name of telescope

    Returns
    -------
    list
        list of targets scheduled on this SSO telescope that  day

    """
    if (telescope == 'Artemis') or (telescope == 'Saint-Ex') or (telescope == 'TS_La_Silla') or (telescope == 'TN_Oukaimeden'):
        telescopes = ['Io', 'Europa', 'Ganymede', 'Callisto']
    else:
        telescopes = ['Io', 'Europa', 'Ganymede', 'Callisto']
        telescopes = np.delete(telescopes, telescopes.index(telescope))
    targets_on_SSO_telescopes = []
    for i in range(len(telescopes)):
        night_block_str =  '/night_blocks_' + telescopes[i] + '_' + str(date) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_SSO_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] + ' on the ' + str(date))

    return targets_on_SSO_telescopes

def SNO_planned_targets(date):
    """ tell which target are scheduled on SNO on a given day

    Parameters
    ----------
    date : date
        date of day in fmt 'yyyy-mm-dd'

    Returns
    -------
    list
        list of targets scheduled on this SNO that  day

    """
    telescopes = ['Artemis', 'Saint-Ex']
    targets_on_SNO_telescopes = []
    for i in range(len(telescopes)):
        night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_SNO_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] + ' on the ' + str(date))
    #print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Targets on SNO: ',targets_on_SNO_telescopes)
    return targets_on_SNO_telescopes

def TS_planned_targets(date):
    """ tell which target are scheduled on TS on a given day

    Parameters
    ----------
    date : date
        date of day in fmt 'yyyy-mm-dd'

    Returns
    -------
    list
        list of targets scheduled on this TS that  day

    """
    telescopes = ['TS_La_Silla']
    targets_on_TS_telescopes = []
    for i in range(len(telescopes)):
        night_block_str =  '/night_blocks_' + telescopes[i] + '_' + str(date) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_TS_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] + ' on the ' + str(date))
    return targets_on_TS_telescopes

def TN_planned_targets(date):
    """ tell which target are scheduled on TN on a given day

    Parameters
    ----------
    date : date
        date of day in fmt 'yyyy-mm-dd'

    Returns
    -------
    list
        list of targets scheduled on this TN that  day

    """
    telescopes = ['TN_Oukaimeden']
    targets_on_TN_telescopes = []
    for i in range(len(telescopes)):
        night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_TN_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] + ' on the ' + str(date))
    return targets_on_TN_telescopes

def target_list_good_coord_format(path_target_list):

    """ Give target corrdinates in ICRS format (used for astropy.coordinates SkyCoord function)

    Parameters
    ----------
    path_target_list: str
        path on your computer toward the target list, by default take the one on the Cambridge server

    Returns
    -------
    targets: astropy.FixedTarget
        targets list with the following format : [<FixedTarget "Sp0002+0115" at SkyCoord (ICRS): (ra, dec) in deg (0.52591667, 1.26003889)>,


    """
    df = pd.read_csv(path_target_list, delimiter=' ')
    target_table_spc = Table.from_pandas(df)
    targets = [FixedTarget(coord=SkyCoord(ra=target_table_spc['RA'][i]* u.degree,dec=target_table_spc['DEC'][i] * u.degree),name=target_table_spc['Sp_ID'][i]) for i in range(len(target_table_spc['RA']))]
    return targets

def charge_observatories(Name):
    """

    Parameters
    ----------
    Name : str
        name of the observatory (ex: 'SSO')

    Returns
    -------
    astroplan.observer
        all info on observatory loaded

    """
    observatories = []
    #Oservatories
    if 'SSO' in str(Name):
        location = EarthLocation.from_geodetic(-70.40300000000002*u.deg, -24.625199999999996*u.deg,2635.0000000009704*u.m)
        observatories.append(Observer(location=location, name="SSO", timezone="UTC"))

    if 'SNO' in str(Name):
        location_SNO = EarthLocation.from_geodetic(-16.50583131*u.deg, 28.2999988*u.deg, 2390*u.m)
        observatories.append(Observer(location=location_SNO, name="SNO", timezone="UTC"))

    if 'Saint-Ex' in str(Name):
        location_saintex = EarthLocation.from_geodetic(-115.48694444444445*u.deg, 31.029166666666665*u.deg, 2829.9999999997976*u.m)
        observatories.append(Observer(location=location_saintex, name="Saint-Ex", timezone="UTC"))

    if 'TS_La_Silla' in str(Name):
        location_TSlasilla = EarthLocation.from_geodetic(-70.73000000000002*u.deg, -29.25666666666666*u.deg, 2346.9999999988418*u.m)
        observatories.append(Observer(location=location_TSlasilla, name="TS_La_Silla", timezone="UTC"))

    if 'TN_Oukaimeden' in str(Name):
        location_TNOuka = EarthLocation.from_geodetic(-7.862263*u.deg,31.20516*u.deg, 2751*u.m)
        observatories.append(Observer(location=location_TNOuka, name="TN_Oukaimeden", timezone="UTC"))

    if 'Munich' in str(Name):
        location_munich= EarthLocation.from_geodetic(48.2*u.deg, -11.6*u.deg, 600*u.m)
        observatories.append(Observer(location=location_munich, name="Munich", timezone="UTC"))

    return observatories

def _generate_24hr_grid(t0, start, end, N, for_deriv=False):
    """ Generate a nearly linearly spaced grid of time durations.
    The midpoints of these grid points will span times from ``t0``+``start``
    to ``t0``+``end``, including the end points, which is useful when taking
    numerical derivatives.
    Parameters
    ----------
    t0 : `~astropy.time.Time`
        Time queried for, grid will be built from or up to this time.
    start : float
        Number of days beFore/after ``t0`` to start the grid.
    end : float
        Number of days beFore/after ``t0`` to end the grid.
    N : int
        Number of grid points to generate
    for_deriv : bool
        Generate time series for taking numerical derivative (modify
        bounds)?
    Returns
    -------
    `~astropy.time.Time`
    """

    if for_deriv:
        time_grid = np.concatenate([[start - 1/(N-1)],
                                    np.linspace(start, end, N)[1:-1],
                                    [end + 1/(N-1)]])*u.day
    else:
        time_grid = np.linspace(start, end, N)*u.day

    # broadcast so grid is first index, and remaining shape of t0
    # falls in later indices. e.g. if t0 is shape (10), time_grid
    # will be shape (N, 10). If t0 is shape (5, 2), time_grid is (N, 5, 2)
    while time_grid.ndim <= t0.ndim:
        time_grid = time_grid[:, np.newaxis]
    # we want to avoid 1D grids since we always want to broadcast against targets
    if time_grid.ndim == 1:
        time_grid = time_grid[:, np.newaxis]
    return t0 + time_grid

def altaz(self, time, target=None, obswl=None, grid_times_targets=False):
    """
    
    :param self:
    :param time:
    :param target:
    :param obswl:
    :param grid_times_targets:
    :return:
    """
    if target is not None:
        time, target = self._preprocess_inputs(time, target, grid_times_targets)

        altaz_frame = AltAz(location=self.location, obstime=time,
                      pressure=self.pressure, obswl=obswl,
                      temperature=self.temperature,
                      relative_humidity=self.relative_humidity)
    if target is None:
      # Return just the frame
      return altaz_frame
    else:
        return target.transform_to(altaz_frame)

def Observability(j,time_range,observatory,targets,constraints):
    """ Give a table with the observability score for each target of targets
    regarding the constraints given and for all ranges of time_range

    Parameters
    ----------
        j : list
            [start end], element of time range (that is to say, month of the year)
        time_range : list
            of astropy.Time range with start and end times
        observatory : astroplan.observer
            observatory chosen
        targets : astropy
            target list on the FixedTarget() format from astroplan
        constraints : astroplan.constraint
            general constraints for a target to be shceduled

    Returns
    -------
        Observability table: astropy.table.Table
             12 columns (target name and each element of time_range, e.g months),
        rows= nb of targets
    """
    targets_observable=[]
    observable=np.argwhere(is_observable(constraints,observatory,targets,time_range=time_range))

    ##WARNING: Need to be replace by a np.where, but I can't find how
    [targets_observable.append(targets[int(obs)]) for obs in observable]

    table_observability=observability_table(constraints, observatory, targets_observable, time_range=time_range, time_grid_resolution = 0.5*u.hour) #calcul un edeuxieme fois is observable
    table_observability.remove_column('always observable')
    table_observability.remove_column('ever observable')
    table_observability.rename_column('fraction of time observable', 'Month' + str(j))

    return table_observability

def reverse_Observability(observatory,targets,constraints,time_ranges):
    """
    Reverse observability table, rows become columns
    Parameters
    ----------
    observatory : str
        observatory chosen
    time_range : list
        List of astropy.Time range with start and end times
    targets : list of astropy.FixedTarget
        target list on the FixedTarget() format from astroplan
    constraints : astroplan.constraints
        general constraints for a target to be shceduled

    Returns
    -------
    reverse_df1 : astropy.table
        observability table with no NaN value (0 instead) inversed with targets as columns
        and elements of time_ranges (months) as rows
    """
    start_fmt = Time(time_ranges[0][0].iso , out_subfmt = 'date').iso
    end_fmt =  Time(time_ranges[len(time_ranges)-1][1].iso , out_subfmt = 'date').iso

    if os.path.exists(path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' +  start_fmt + '_' + end_fmt + '_'  + str(len(targets)) + '.csv'):
        name_file = path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' +  start_fmt + '_' + end_fmt + '_'  + str(len(targets)) +  '.csv'
        reverse_df1 = pd.read_csv(name_file, delimiter = ',')
        return reverse_df1

    else:
        tables_observability=list(map((lambda x: Observability(x,time_ranges[x],observatory,targets,constraints)), range(0,len(time_ranges))))
        df = list(map((lambda x: pd.DataFrame(tables_observability[x].to_pandas())),range(0,len(time_ranges))))
        a = reduce(functools.partial(pd.merge,how='outer', on='target name'), df)
        df = a.replace(to_replace=float('NaN'), value=0.0)
        df1 = df.set_index('target name')
        reverse_df1 = df1.T
        reverse_df1.to_csv(path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' +  start_fmt + '_' + end_fmt + '_'  + str(len(targets)) + '.csv', sep= ',')
        return reverse_df1

def month_option(target_name,reverse_df1):
    """ create a list of the best month for oservation for each target

    Parameters
    ----------
    target_name : str
        name of the target
    reverse_df1 : astropy.table
        observability table with no NaN value (0 instead) inversed with targets as columns
        and elements of time_ranges (months) as rows

    Returns
    -------
    month : list
        a list with the best month to observe the target
    month_2nd_option : list
        same but for the second best month
    months_3rd_option : list
        same but for the third best month
    months_4th_option : list
        same but for the fourth best month
    months_5th_option : list
        same but for the fiveth best month


    Remarks
    -------
        the 2nd, 3rd etc choices are here in case the target list is not long enough to give
        solutions for all the telescopes each months, allows to avoid blancks in observations
    """

    try:
        months = reverse_df1[str(target_name)].idxmax()
        months_2nd_option = reverse_df1[str(target_name)].nlargest(2,keep='first').index[1]
        months_3rd_option = reverse_df1[str(target_name)].nlargest(3,keep='first').index[2]
        months_4th_option = reverse_df1[str(target_name)].nlargest(4,keep='first').index[3]
        months_5th_option= reverse_df1[str(target_name)].nlargest(5,keep='first').index[4]

    except KeyError:
        months = 0
        months_2nd_option = 0
        months_3rd_option = 0
        months_4th_option = 0
        months_5th_option = 0
    return [months, months_2nd_option , months_3rd_option,months_4th_option, months_5th_option]

def save_schedule(input_file,nb_observatory,save,over_write,date_range,telescope):
    """ save schedules in destination

    Parameters
    ----------
    input_file : file
        input file for SPOCKLT
    nb_observatory : str
        name of the observatory
    save : bool
        True if want to save, False if not
    over_write : bool
        True if want overwrite, False if not
    date_range : list of date
         list with 2 elements, start date  and end date
    telescope : str
        name of telescope for which you which to save the schedules

    Returns
    -------
    message

    """
    if input_file is None:
        telescope = telescope
        date_range = date_range
        date_range_in_days = int((date_range[1] - date_range[0]).value)
    else:
        with open(input_file, "r") as f:
            Inputs = yaml.load(f, Loader=yaml.FullLoader)
            df = pd.DataFrame.from_dict(Inputs['observatories'])
            observatory = charge_observatories(df[nb_observatory]['name'])[0]
            date_range = Time(Inputs['date_range'])
            date_range_in_days = int((date_range[1]- date_range[0]).value)
            telescope = df[nb_observatory]['telescopes'][0]
    for i in range(0,date_range_in_days):
        day = date_range[0] + i
        if save:
            source = path_spock + '/' + 'night_blocks_propositions/' +'night_blocks_' + telescope + '_' +  day.tt.datetime.strftime("%Y-%m-%d") + '.txt'
            destination = path_spock + '/DATABASE/' + telescope + '/'
            destination_2 = path_spock + '/DATABASE/' + telescope + '/' + 'Archive_night_blocks/'
            if over_write:
                dest = shutil.copy(source, destination)
                dest2 = shutil.copy(source, destination_2)
                print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been over-written to ' + '\"' +  destination + '\"' )
                print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been copied to ' + '\"' + destination_2 + '\"')
            if not over_write:
                try:
                    dest = shutil.move(source, destination)
                    #dest2 = shutil.move(source, destination_2)
                    print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' +  source + '\"' +  ' has been copied to ' + '\"' + destination + '\"' )
                    #print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been copied to ' + '\"' + destination_2 + '\"')
                except shutil.Error:
                    print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + destination + 'night_blocks_' + telescope + '_' +  day.tt.datetime.strftime("%Y-%m-%d") + '.txt' + '\"' +  ' already exists')
        if not save:
            print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + ' Those plans have not been saved')

def make_plans(day, nb_days, telescope):
    """ make plans for telescope for a certain number of day from a start day

    Parameters
    ----------
    day : int
        day
    nb_days : int
        number of days
    telescope : string
        name telescope

    Returns
    -------


    """

    make_np(day, nb_days, telescope)

def upload_plans(day, nb_days, telescope):
    """ upload plans to DATABASE

    Parameters
    ----------
    day : date
        date in fmt 'yyyy-mm-dd'
    nb_days : int
        number of days
    telescope : str
        name of telescope

    Returns
    -------

    """
    if telescope.find('Callisto') is not -1:
        SPOCKunp.upload_np_calli(day, nb_days)
    if telescope.find('Ganymede') is not -1:
        SPOCKunp.upload_np_gany(day, nb_days)
    if telescope.find('Io') is not -1:
        SPOCKunp.upload_np_io(day, nb_days)
    if telescope.find('Europa') is not -1:
        SPOCKunp.upload_np_euro(day, nb_days)
    if telescope.find('Artemis') is not -1:
        SPOCKunp.upload_np_artemis(day, nb_days)
    if telescope.find('TS_La_Silla') is not -1:
        SPOCKunp.upload_np_ts(day, nb_days)
    if telescope.find('TN_Oukaimeden') is not -1:
        SPOCKunp.upload_np_tn(day, nb_days)
    if telescope.find('Saint-Ex') is not -1:
        SPOCKunp.upload_np_saint_ex(day, nb_days)


    # ------------------- update archive date by date plans folder  ------------------

    path_database = os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/', telescope,'schedule')
    path_gant_chart = os.path.join(path_spock + '/SPOCK_Figures/Preview_schedule.html')
    path_gant_chart_masterfile = os.path.join('/Users/elsaducrot/spock_2/SPOCK_Files/spock_stats_masterfile.csv')
    path_database_home = os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/Preview_schedule.html')
    path_database_home_masterfile = os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/spock_stats_masterfile.csv')
    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path local \'Gant chart\' = ', path_gant_chart)
    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path database \'Gant chart\' = ',  path_database_home)
    subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart, path_database_home])
    subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart_masterfile, path_database_home_masterfile])


class Schedules:
    """
    Class to Make schedules for the target list, observatory, date_range and startegy indicated

    """


    def __init__(self):


        self.Altitude_constraint = 25
        self.constraints = None
        self.date_range = None  # date_range
        self.dur_obs_set_target = None
        self.dur_obs_rise_target =  None
        self.dur_obs_both_target = None
        self.duration_segments = None  # duration_segments
        self.first_target = None
        self.first_target_by_day = []
        self.idx_first_target = None
        self.idx_first_target_by_day = []
        self.idx_second_target = None
        self.idx_second_target_by_day = []
        self.index_prio = None
        self.index_prio_by_day = []
        self.idx_planned_SSO = None
        self.idx_SSO_in_planned = None
        self.moon_and_visibility_constraint_table = None
        self.nb_segments =  None #nb_segments
        self.night_block = []
        self.night_block_by_day = []
        self.observatory = None  # observatory
        self.observability_table_day = None
        self.planned_targets = []
        self.priority = None
        self.priority_by_day = []
        self.priority_ranked = None
        self.priority_ranked_by_day = []
        self.reverse_df1 = None
        self.second_target = None
        self.second_target_by_day = []
        self.strategy = None
        self.targets = None
        self.target_list = None
        self.target_table_spc = []
        self.telescopes = []
        self.telescope = []


    @property
    def idx_rise_targets_sorted(self):
        """ index for rise targets sorted
        
        Returns
        -------
        list of int
            index of rise targets sorted

        """
        idx_rise_targets=(self.priority_ranked['set or rise']=='rise')
        idx_rise_targets_sorted=self.index_prio[idx_rise_targets]
        return idx_rise_targets_sorted

    @property
    def idx_set_targets_sorted(self):
        """ index for set targets sorted

        Returns
        -------
        list of int
            index of set targets sorted
        """
        idx_set_targets=(self.priority_ranked['set or rise']=='set')
        idx_set_targets_sorted=self.index_prio[idx_set_targets]
        return idx_set_targets_sorted

    @property
    def months_obs(self):
        """ month of obs

        Returns
        -------
        int
            month number (between 0 and 11)

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        for i,t in enumerate(self.time_ranges):
            if (datetime.strptime(t[0].value,date_format) <= datetime.strptime(self.date_range[0].value,date_format) <= datetime.strptime(t[1].value,date_format)) and\
                    (datetime.strptime(t[0].value,date_format) <= datetime.strptime(self.date_range[1].value,date_format) <= datetime.strptime(t[1].value,date_format)):
                return  i
            if (datetime.strptime(t[0].value,date_format) <= datetime.strptime(self.date_range[0].value,date_format) <= datetime.strptime(t[1].value,date_format)) and\
                    (datetime.strptime(t[1].value,date_format) <= datetime.strptime(self.date_range[1].value,date_format)) :
                if i < (len(self.time_ranges) - 1):
                    return i+1
                if i == (len(self.time_ranges) - 1):
                    return i

    @property
    def date_range_in_days(self):
        """ number of days in date range

        Returns
        -------
        int
            number of day between date start and date end

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        date_start = datetime.strptime(self.date_range[0].value, date_format)
        date_end = datetime.strptime(self.date_range[1].value, date_format)
        date_range_in_days = (date_end - date_start).days
        return date_range_in_days

    @property
    def nb_hours_threshold(self):
        """ number of hours to reach

        Returns
        -------
        list
            list as long as the target list

        """
        nb_hours_threshold = [100]* len(self.target_table_spc)
        return nb_hours_threshold

    @property
    def date_ranges_day_by_day(self):
        """ date range day by day

        Returns
        -------
        list
            list of date ranges

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        d2 = datetime.strptime(self.date_range[1].value, date_format)
        i = 0
        t = datetime.strptime(self.date_range[0].value, date_format)
        u = t
        date_ranges_day_by_day = []
        while t < d2:
            d = timedelta(days=1)
            t = u + d * i
            date_ranges_day_by_day.append(Time(t))
            i += 1
        return date_ranges_day_by_day

    @property
    def nb_hours_observed(self):
        """ nb hours observed for each target

        Returns
        -------
        list
            nb hours observed

        """
        nb_hours_observed = self.target_table_spc['nb_hours_surved']
        return nb_hours_observed

    @property
    def time_ranges(self):
        year = self.date_range[0].tt.datetime.strftime("%Y")
        time_ranges = [Time([year + '-01-01 12:00:00', year + '-01-31 12:00:00']),
                            Time([year + '-02-01 12:00:00', year + '-02-28 12:00:00']),
                            Time([year + '-03-01 15:00:00', year + '-03-31 15:00:00']),
                            Time([year + '-04-01 15:00:00', year + '-04-30 15:00:00']),
                            Time([year + '-05-01 15:00:00', year + '-05-31 15:00:00']),
                            Time([year + '-06-01 15:00:00', year + '-06-30 15:00:00']),
                            Time([year + '-07-01 12:00:00', year + '-07-31 12:00:00']),
                            Time([year + '-08-01 12:00:00', year + '-08-31 12:00:00']),
                            Time([year + '-09-01 12:00:00', year + '-09-30 12:00:00']),
                            Time([year + '-10-01 12:00:00', year + '-10-31 12:00:00']),
                            Time([year + '-11-01 12:00:00', year + '-11-30 12:00:00']),
                            Time([year + '-12-01 12:00:00', year + '-12-31 12:00:00'])]
        return time_ranges

    def exposure_time_table(self,day):
        """ generate an exposure time table as the form of a file untitled

        Parameters
        ----------
        day : date
            date in fmt 'yyyy-mm-dd'

        Returns
        -------
        file
            file with most appropriate exposure time for each target "exposure_time_table.csv"

        """
        if day is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
        sso_texp = np.zeros(len(self.target_table_spc))
        sno_texp = np.zeros(len(self.target_table_spc))
        saintex_texp = np.zeros(len(self.target_table_spc))
        ts_texp = np.zeros(len(self.target_table_spc))
        tn_texp = np.zeros(len(self.target_table_spc))
        for i in range(len(self.target_table_spc)):
            print(i,self.target_table_spc['Sp_ID'][i])
            sso_texp[i] = self.exposure_time_for_table('SSO', day, i)
            sno_texp[i] = self.exposure_time_for_table( 'SNO', day, i)
            saintex_texp[i] = self.exposure_time_for_table('Saint-Ex', day, i)
            ts_texp[i] = self.exposure_time_for_table('TS', day, i)
            tn_texp[i] = self.exposure_time_for_table('TN', day, i)

            df = pd.DataFrame({'Sp_ID':self.target_table_spc['Sp_ID'],\
                               'SSO_texp':sso_texp,'SNO_texp':sno_texp,'Saintex_texp':saintex_texp,\
                               'TS_texp':ts_texp,'TN_texp':tn_texp,})
            df.to_csv(path_spock + '/SPOCK_files/exposure_time_table.csv',sep=',',index=False)

    def idx_SSO_observed_targets(self):
        df_cambridge = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        df_cambridge['Target'] = [x.strip().replace('SP', 'Sp') for x in df_cambridge['Target']]
        server_in_targetlist, targetlist_in_server = index_list1_list2(df_cambridge['Target'], self.target_table_spc['Sp_ID'])
        return server_in_targetlist

    def idx_SNO_observed_targets(self):
        df_artemis = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        artemis_in_targetlist, targetlist_in_artemis = index_list1_list2(df_artemis['Target'], self.target_table_spc['Sp_ID'])
        return artemis_in_targetlist

    def idx_SaintEx_observed_targets(self):
        df_saintex = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', delimiter=',')
        saintex_in_targetlist, targetlist_in_saintex = index_list1_list2(df_saintex['Target'], self.target_table_spc['Sp_ID'])
        return saintex_in_targetlist

    def idx_trappist_observed_targets(self):
        df_trappist = pd.read_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt', delimiter=',')
        trappist_in_targetlist, targetlist_in_trappist = index_list1_list2(df_trappist['Target'], self.target_table_spc['Sp_ID'])
        return trappist_in_targetlist

    def load_parameters(self,nb_observatory=None,input_file=None,date_range=None,obs_name=None):
        if input_file == None:
            self.telescopes = ['Io','Europa','Ganymede','Callisto']
            self.strategy = 'continuous'
            self.target_list = path_spock + '/target_lists/speculoos_target_list_v6.txt'
            self.constraints = [AtNightConstraint.twilight_nautical()]
            df = pd.read_csv(self.target_list, delimiter=' ')
            self.observatory = charge_observatories(obs_name)[0]
            self.target_table_spc = Table.from_pandas(df)
            self.targets = target_list_good_coord_format(self.target_list)

            last_mod = time.ctime(os.path.getmtime(self.target_list))
            now = datetime.now()
            current_time = now.strftime("%Y-%m-%d %H:%M:%S")
            time_since_last_update = (Time(datetime.strptime(last_mod, "%a %b %d %H:%M:%S %Y"), format='datetime') - \
                                      Time(datetime.strptime(current_time, "%Y-%m-%d %H:%M:%S"),
                                           format='datetime')).value * 24
            # self.update_nb_hours_all()
            if abs(time_since_last_update) > 24:  # in hours
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Updating the number of hours observed')
                self.update_nb_hours_all()
            self.date_range = Time(date_range)
            if self.date_range[1] <= self.date_range[0]:
                sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' end date inferior to start date')

        else:
            with open(input_file, "r") as f:
                Inputs = yaml.load(f, Loader=yaml.FullLoader)
                self.target_list = path_spock + '/target_lists/speculoos_target_list_v6.txt'
                df = pd.read_csv(self.target_list, delimiter=' ')
                self.target_table_spc = Table.from_pandas(df)
                df = pd.DataFrame.from_dict(Inputs['observatories'])
                self.observatory = charge_observatories(df[nb_observatory]['name'])[0]
                self.telescopes = df[nb_observatory]['telescopes']
                self.telescope = self.telescopes[0]
                try:
                    self.date_range = Time(Inputs['date_range']) #,Time(Inputs['date_range'][1])]
                except ValueError:
                    sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' Wrong date format, date must be %y-%m-%d HH:MM:SS.sss')
                if self.date_range[1] <= self.date_range[0]:
                    sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' end date inferior to start date')
                self.strategy = Inputs['strategy']
                self.duration_segments = Inputs['duration_segments']
                self.nb_segments = Inputs['nb_segments']
                self.constraints = [AtNightConstraint.twilight_nautical()]
                df = pd.read_csv(self.target_list, delimiter=' ')
                self.target_table_spc = Table.from_pandas(df)
                self.targets = target_list_good_coord_format(self.target_list)

                last_mod = time.ctime(os.path.getmtime(self.target_list))
                now = datetime.now()
                current_time = now.strftime("%Y-%m-%d %H:%M:%S")
                time_since_last_update = (Time(datetime.strptime(last_mod, "%a %b %d %H:%M:%S %Y"),format='datetime') - \
                                         Time(datetime.strptime(current_time, "%Y-%m-%d %H:%M:%S"), format='datetime')).value * 24
                #self.update_nb_hours_all()
                if abs(time_since_last_update) > 24: # in hours
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Updating the number of hours observed')
                    self.update_nb_hours_all()

    def update_nb_hours_all(self,user =user_portal, password = pwd_portal):
        # *********** TRAPPIST ***********
        #self.get_hours_files_TRAPPIST()

        # *********** SSO & SNO ***********
        self.update_telescope_from_server() # get hours SSO
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyTotal"
        target_list = pd.read_csv(self.target_list, delimiter=' ')
        target_list['telescope'] =  ['None'] * len(target_list['telescope'])
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyTotal.txt', 'wb').write(resp.content)
        df =  pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/SurveyTotal.txt',sep=' ',index=False)
        df_camserver = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        #df_camserver['Target'][9] = 'Sp0004-2058'
        df_camserver['Target'] = [x.strip().replace('SP', 'Sp') for x in df_camserver['Target']]

        # *********** Saint-Ex ***********
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt',delimiter=',')
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt',sep=',',index=False)

        # Read files
        df_SaintEx = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', delimiter=',')
        df_TRAPPIST = pd.read_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt', delimiter=',')
        df_camserver_telescope = pd.read_csv(path_spock + '/survey_hours/SurveyByTelescope.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        df_camserver_telescope['Target'] = [x.strip().replace('SP', 'Sp') for x in df_camserver_telescope['Target']]
        for i in range(len(target_list)):
            idxs = np.where((df_camserver_telescope['Target'] == target_list['Sp_ID'][i]))[0]
            target_list['telescope'][i] = list(df_camserver_telescope['Telescope'][idxs])

        SaintEx_in_targetlist, targetlist_in_SaintEx = index_list1_list2(df_SaintEx['Target'], target_list['Sp_ID'])
        camserver_in_targetlist, targetlist_in_camserver = index_list1_list2(df_camserver['Target'], target_list['Sp_ID'])
        TRAPPIST_in_targetlist, targetlist_in_TRAPPIST = index_list1_list2(df_TRAPPIST['Target'],target_list['Sp_ID'])

        with alive_bar(len(target_list)) as bar:
            for i in range(len(target_list)):
                bar()
                time.sleep(0.001)

                if np.any((np.asarray(camserver_in_targetlist) == i)): #SSO or SNO
                    idx_camserver = np.argwhere((np.asanyarray(camserver_in_targetlist) == i))[0][0]
                    target_list['nb_hours_surved'][i] = df_camserver['Hours'][targetlist_in_camserver[idx_camserver]]

                    if np.any((np.asarray(SaintEx_in_targetlist) == i)) \
                            and np.any((np.asarray(TRAPPIST_in_targetlist) == i)): #Saint-Ex and TRAPPIST
                        idx_SaintEx = np.argwhere((np.asanyarray(SaintEx_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += df_SaintEx[' Observation Hours '][targetlist_in_SaintEx[idx_SaintEx]]
                        target_list['telescope'][i].append('Saint-Ex')
                        # if np.any((np.asarray(TRAPPIST_in_targetlist) == i)):
                        idx_TRAPPIST = np.argwhere((np.asanyarray(TRAPPIST_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += df_TRAPPIST[' Observation Hours '][targetlist_in_TRAPPIST[idx_TRAPPIST]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(TRAPPIST_in_targetlist) == i)) \
                            and np.all((np.asarray(SaintEx_in_targetlist) != i)): #TRAPPIST but not Saint-Ex
                        idx_TRAPPIST = np.argwhere((np.asanyarray(TRAPPIST_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += df_TRAPPIST[' Observation Hours '][targetlist_in_TRAPPIST[idx_TRAPPIST]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(SaintEx_in_targetlist) == i)) \
                            and np.all((np.asarray(TRAPPIST_in_targetlist) != i)): #Saint-Ex but not TRAPPIST
                        idx_SaintEx = np.argwhere((np.asanyarray(SaintEx_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += df_SaintEx[' Observation Hours '][targetlist_in_SaintEx[idx_SaintEx]]
                        target_list['telescope'][i].append('Saint-Ex')

                if np.all((np.asarray(camserver_in_targetlist) != i)):  # Neither SSO nor SNO

                    if np.any((np.asarray(SaintEx_in_targetlist) == i)) \
                            and np.any((np.asarray(TRAPPIST_in_targetlist) == i)) : # Saint-Ex and TRAPPIST
                        idx_SaintEx = np.argwhere((np.asanyarray(SaintEx_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = df_SaintEx[' Observation Hours '][targetlist_in_SaintEx[idx_SaintEx]]
                        target_list['telescope'][i].append('Saint-Ex')
                        idx_TRAPPIST = np.argwhere((np.asanyarray(TRAPPIST_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += df_TRAPPIST[' Observation Hours '][targetlist_in_TRAPPIST[idx_TRAPPIST]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(TRAPPIST_in_targetlist) == i)) \
                            and np.all((np.asarray(SaintEx_in_targetlist) != i)): #TRAPPIST but not Saint-Ex
                        idx_TRAPPIST = np.argwhere((np.asanyarray(TRAPPIST_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = df_TRAPPIST[' Observation Hours '][targetlist_in_TRAPPIST[idx_TRAPPIST]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(SaintEx_in_targetlist) == i)) \
                            and np.all((np.asarray(TRAPPIST_in_targetlist) != i)): #Saint-Ex but not TRAPPIST
                        idx_SaintEx = np.argwhere((np.asanyarray(SaintEx_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = df_SaintEx[' Observation Hours '][targetlist_in_SaintEx[idx_SaintEx]]
                        target_list['telescope'][i].append('Saint-Ex')

        target_list.to_csv(self.target_list, sep=' ', index=False)
        target_list.to_csv(path_spock + '/target_lists/speculoos_target_list_v6_sep_coma.csv', sep=',')
        subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/target_lists/speculoos_target_list_v6.txt',
                          'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/target_lists/'])

    def get_hours_files_TRAPPIST(self):
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        #creds = ServiceAccountCredentials.from_json_keyfile_name('client_secret2.json', scope)
        #client = gspread.authorize(creds)
        sheet = client.open('Hours_observation_TS_TN').sheet1
        # Extract and print all of the values
        #list_of_hashes = sheet.get_all_records()
        col2 = sheet.col_values(2)
        col5 = sheet.col_values(6)
        col5 = col5[13:]
        col5 = [float(col5[i].replace(",", ".")) for i in range(0, len(col5))]
        target_observed_TSTN = pd.Series(col2[13:])
        hours_observed_TSTN = pd.Series(col5)
        telescopes = ['TRAPPIST'] * len(target_observed_TSTN)
        df_google_doc = pd.DataFrame({'Target': target_observed_TSTN, ' Observation Hours ': hours_observed_TSTN,'telescope':telescopes})
        df_google_doc.to_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt',sep=',',index=False)
        subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt',
                          'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/survey_hours/'])

    def update_nb_hours_SNO(self):
        get_hours_files_SNO()
        target_list = pd.read_csv(self.target_list, delimiter=' ')
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        SNO_in_targetlist, targetlist_in_SNO = index_list1_list2(df['Target'], target_list['Sp_ID'])
        target_list['nb_hours_surved'][SNO_in_targetlist] = df[' Observation Hours '][targetlist_in_SNO]
        #target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_telescope_SNO(self):
        try:
            get_hours_files_SNO()
        except TimeoutError:
            print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' Are on the Liege  VPN ?')
        target_list = pd.read_csv(self.target_list, delimiter=' ')
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        SNO_in_targetlist, targetlist_in_SNO = index_list1_list2(df['Target'], target_list['Sp_ID'])
        df_tel = ['Artemis'] * len(df['Target'][targetlist_in_SNO])
        target_list['telescope'][SNO_in_targetlist] = df_tel
        target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_nb_hours_from_server(self,user = 'educrot',password = pwd_portal):
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyTotal"
        target_list = pd.read_csv(self.target_list, delimiter=' ')
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyTotal.txt', 'wb').write(resp.content)
        df = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        df['Target'][9] = 'Sp0004-2058'
        df['Target'] = [x.strip().replace('SP', 'Sp') for x in df['Target']]
        server_in_targetlist, targetlist_in_server = index_list1_list2(df['Target'], target_list['Sp_ID'])
        target_list['nb_hours_surved'][server_in_targetlist] = df['Hours'][targetlist_in_server]
        #target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_telescope_from_server(self,user =  'educrot',password = pwd_portal):
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyByTelescope"
        target_list = pd.read_csv(self.target_list, delimiter=' ')
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyByTelescope.txt', 'wb').write(resp.content)
        df = pd.read_csv(path_spock + '/survey_hours/SurveyByTelescope.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/SurveyByTelescope.txt',sep=' ',index=False)
        #
        #df = pd.read_csv('SurveyByTelescope.txt', delimiter=' ', skipinitialspace=True, error_bad_lines=False)
        #df['Target'] = [x.strip().replace('SP', 'Sp') for x in df['Target']]
        #server_in_targetlist, targetlist_in_server = index_list1_list2(df['Target'], target_list['Sp_ID'])
        #target_list['telescope'][server_in_targetlist] = df['Telescope'][targetlist_in_server]

        #for i in range(len(target_list)):
        #    idxs = np.where((df['Target'] == target_list['Sp_ID'][i]))[0]
        #    target_list['telescope'][i] == list(df['Telescope'][idxs])
        #    print(target_list['telescope'][i])

        #target_list.to_csv(self.target_list,sep=' ',index=False)

    def make_schedule(self, Altitude_constraint = None, Moon_constraint = None):
        import time
        start = time.time()
        # for telescope in self.telescopes:
        self.nb_hours = np.zeros((len(self.target_table_spc['nb_hours_surved']),len(self.target_table_spc['nb_hours_surved'])))
        self.nb_hours[0,:]  = self.target_table_spc['nb_hours_surved']
        self.nb_hours[1,:] = self.target_table_spc['nb_hours_surved']
        self.targets = target_list_good_coord_format(self.target_list)

        observed_targets_SSO = []
        observed_targets_SNO = []
        observed_targets_TS = []
        observed_targets_TN = []
        for i in tqdm(range(self.date_range_in_days),desc="Updating hours of obs "):
            pass
            date = self.date_range[0] + i
            day_fmt = Time(date.iso, out_subfmt='date').iso
            observed_targets_SSO += SSO_planned_targets(day_fmt,self.telescope)
            observed_targets_SNO += SNO_planned_targets(day_fmt)
            observed_targets_TS += TS_planned_targets(day_fmt)
            observed_targets_TN += TN_planned_targets(day_fmt)
        observed_targets_SSO = list(set(observed_targets_SSO))
        observed_targets_SNO = list(set(observed_targets_SNO))
        observed_targets_TS = list(set(observed_targets_TS))
        observed_targets_TN = list(set(observed_targets_TN))

        self.idx_planned_SSO, self.idx_SSO_in_planned = index_list1_list2(observed_targets_SSO,self.target_table_spc['Sp_ID'])
        self.idx_planned_SNO, self.idx_SNO_in_planned = index_list1_list2(observed_targets_SNO,self.target_table_spc['Sp_ID'])
        self.idx_planned_TS, self.idx_TS_in_planned = index_list1_list2(observed_targets_TS,self.target_table_spc['Sp_ID'])
        self.idx_planned_TN, self.idx_TN_in_planned = index_list1_list2(observed_targets_TN,self.target_table_spc['Sp_ID'])

        if Altitude_constraint:
            self.constraints.append(AltitudeConstraint(min=float(Altitude_constraint)*u.deg))
        if Moon_constraint:
            self.constraints.append(MoonSeparationConstraint(min=float(Moon_constraint)*u.deg))

        if str(self.strategy) == 'alternative':
            self.is_moon_and_visibility_constraint()

        if str(self.strategy) == 'continuous':
            self.reverse_df1 = reverse_Observability(self.observatory,self.targets,self.constraints,self.time_ranges)
            end = time.time()
            #print('reverse_df1',end - start)

            for t in tqdm(range(0,self.date_range_in_days),desc="Scheduling "):
                pass
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' day is : ',Time(self.date_range[0] + t).iso)
                day = self.date_ranges_day_by_day[t]
                product_table_priority_prio = self.table_priority_prio(day)
                self.index_prio = product_table_priority_prio[0]
                self.priority = product_table_priority_prio[1]
                self.idx_targets(t)
                self.moon_and_visibility_constraint_table = self.is_moon_and_visibility_constraint(day)
                self.priority_by_day.append(self.priority)
                self.index_prio_by_day.append(self.index_prio)
                self.priority_ranked_by_day.append(self.priority_ranked)
                end = time.time()
                #print('Ranked prio', end - start)

                if self.is_constraints_met_first_target(t):
                    self.first_target = self.priority[self.idx_first_target]
                    self.first_target_by_day.append(self.first_target)
                    self.idx_first_target_by_day.append(self.idx_first_target)

                if not self.is_constraints_met_first_target(t):
                    print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' impossible to find the first target that respects the constraints')

                if self.idx_second_target != None and self.is_constraints_met_second_target(t):
                    self.second_target = self.priority[self.idx_second_target]
                    self.second_target_by_day.append(self.second_target)
                    self.idx_second_target_by_day.append(self.idx_second_target)

                if self.idx_second_target is None:
                    print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' no second target')

                self.night_block = self.schedule_blocks(day)
                self.night_block_by_day.append(self.night_block)
                self.make_night_block(day)
                self.make_plan_file(day)

        if str(self.strategy) == 'segmented':
            print()

    def idx_targets(self, t):
        """
            Give the index of the first and second targets as well as the
            corresponding row in the priority table

        Parameters
        ----------

        Returns
        -------
            idx_first_target: int
                index first target in target list
            first_target: int
                row number idx_first_target in priority table
            idx_second_target: int
                index second target in target list
            second_target: row number idx_second_target in priority table

        """
        idx_init_first = -1
        self.idx_first_target = self.index_prio[idx_init_first]
        self.first_target=self.priority[self.idx_first_target]
        dt_1day=Time('2018-01-02 00:00:00',scale='tcg')-Time('2018-01-01 00:00:00',scale='tcg') #1 day

        # if self.target_table_spc['texp_spc'][self.idx_first_target] == 0:
        #     self.target_table_spc['texp_spc'][self.idx_first_target]= self.exposure_time(self.date_ranges_day_by_day[t],self.idx_first_target)

        #if (self.telescope  == 'Artemis') or (self.telescope  == 'Saint-Ex'):
            #while (self.target_table_spc['texp_spc'][self.idx_first_target] > 150) or ('Io' in self.target_table_spc['telescope'][self.idx_first_target]) or ('Europa' in self.target_table_spc['telescope'][self.idx_first_target]) or \
            #        ('Ganymede' in self.target_table_spc['telescope'][self.idx_first_target]) or ('Callisto' in self.target_table_spc['telescope'][self.idx_first_target]) or ('TRAPPIST' in self.target_table_spc['telescope'][self.idx_first_target]):
            #    idx_init_first -= 1
            #    self.idx_first_target = self.index_prio[idx_init_first]
            #    self.first_target = self.priority[self.idx_first_target]
            #    self.target_table_spc['texp_spc'][self.idx_first_target] = self.exposure_time(self.date_ranges_day_by_day[t],self.idx_first_target)

        if (self.telescope == 'Io') or (self.telescope == 'Europa') or (self.telescope == 'Ganymede') or (self.telescope == 'Callisto'):
            other_SSO = np.delete(self.telescopes, self.telescopes.index(self.telescope))
            while (self.target_table_spc['texp_spc'][self.idx_first_target] > 150) or np.any([other_SSO[j] in self.target_table_spc['telescope'][self.idx_first_target] for j in range(len(list(other_SSO)))]) or\
                    ('Artemis' in self.target_table_spc['telescope'][self.idx_first_target])  or ('Saint-Ex' in self.target_table_spc['telescope'][self.idx_first_target]) \
                    or ('TRAPPIST' in self.target_table_spc['telescope'][self.idx_first_target]):
                idx_init_first -= 1
                self.idx_first_target = self.index_prio[idx_init_first]
                self.first_target = self.priority[self.idx_first_target]
                # if self.target_table_spc['texp_spc'][self.idx_first_target] == 0:
                #     self.target_table_spc['texp_spc'][self.idx_first_target] = self.exposure_time(self.date_ranges_day_by_day[t],self.idx_first_target)

        if (self.telescope == 'TS_La_Silla') or (self.telescope == 'TN_Oukaimeden'):
            while (self.target_table_spc['texp_spc'][self.idx_first_target] > 150) or np.any([self.telescopes[j] in self.target_table_spc['telescope'][self.idx_first_target] for j in range(len(list(self.telescopes)))])  \
                    or ('Artemis' in self.target_table_spc['telescope'][self.idx_first_target]) or ('Saint-Ex' in self.target_table_spc['telescope'][self.idx_first_target]) :
                print('test ')
                idx_init_first -= 1
                self.idx_first_target = self.index_prio[idx_init_first]
                self.first_target = self.priority[self.idx_first_target]

        for i in range(1,abs(idx_init_first)+len(self.index_prio)):
            # print(i, self.targets[self.index_prio[-(i)]])
            rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t,self.targets[self.idx_first_target],which='next',horizon=24*u.deg)
            set_first_target = self.observatory.target_set_time(self.date_range[0]+t,self.targets[self.idx_first_target],which='next',horizon=24*u.deg)
            if self.observatory.target_set_time(self.date_range[0] + t, self.targets[self.idx_first_target], which='next',
                                             horizon=24 * u.deg) < self.observatory.target_rise_time(self.date_range[0] + t,self.targets[self.idx_first_target],which='nearest',horizon=24*u.deg):
                set_first_target = self.observatory.target_set_time(self.date_range[0]+t+1,self.targets[self.idx_first_target],which='next',horizon=24*u.deg)
                if self.observatory.target_rise_time(self.date_range[0] + t, self.targets[self.idx_first_target],
                        which='next',horizon=24 * u.deg) < self.observatory.target_rise_time(self.date_range[0] + t,
                                            self.targets[self.idx_first_target], which='nearest', horizon=24 * u.deg):
                    rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t + 1,self.targets[self.idx_first_target], which='next',
                                                                          horizon=24 * u.deg)

            if self.telescope == 'Saint-Ex':
                set_target = self.observatory.target_set_time(self.date_range[0] + t,self.targets[self.index_prio[-i]],which='next',horizon=24*u.deg)
                rise_target = self.observatory.target_rise_time(self.date_range[0] + t,self.targets[self.index_prio[-i]],which='next',horizon=24*u.deg)
                if set_target > (self.date_range[0] + t +1):
                    set_target = self.observatory.target_set_time(self.date_range[0]+ t,self.targets[self.index_prio[-i]], which='nearest',horizon=24 * u.deg)
                if rise_target < (self.date_range[0] + t ):
                    rise_target = self.observatory.target_set_time(self.date_range[0]+t+1,self.targets[self.index_prio[-i]], which='nearest',horizon=24 * u.deg)
            else:
                set_target = self.observatory.target_set_time(self.date_range[0] + t,self.targets[self.index_prio[-i]],which='next',horizon=24*u.deg)
                rise_target = self.observatory.target_rise_time(self.date_range[0] + t,self.targets[self.index_prio[-i]],which='next',horizon=24*u.deg)

            if self.first_target['set or rise'] == 'rise':
                if self.priority['set or rise'][self.index_prio[-i]]=='set':
                    if (rise_target < self.observatory.twilight_evening_nautical(self.date_range[0]+t,which='next')) and \
                            (set_target > rise_first_target):
                        self.idx_second_target=self.index_prio[-i]
                        self.second_target = self.priority[self.idx_second_target]
                        break

                    else:
                        self.second_target = None
                        self.idx_second_target = None

                if self.priority['set or rise'][self.index_prio[-i]]=='both':
                    if (rise_target < self.observatory.twilight_evening_nautical(self.date_range[0]+t,which='nearest')) and \
                            (set_target > rise_first_target):
                        self.idx_second_target=self.index_prio[-i]
                        self.second_target = self.priority[self.idx_second_target]
                        break


            if self.first_target['set or rise'] == 'set':
                if self.priority['set or rise'][self.index_prio[-i]]=='rise':
                    if (set_target > self.observatory.twilight_morning_nautical(self.date_range[0]+t,which='next')) and \
                            (rise_target < set_first_target):
                        self.idx_second_target = self.index_prio[-i]
                        self.second_target = self.priority[self.idx_second_target]
                        break

                    else:
                        self.second_target = None
                        self.idx_second_target = None

                if self.priority['set or rise'][self.index_prio[-i]]=='both':
                    if (set_target > self.observatory.twilight_morning_nautical(self.date_range[0]+t,which='next')) and \
                            (rise_target < set_first_target):
                        self.idx_second_target=self.index_prio[-i]
                        self.second_target = self.priority[self.idx_second_target]
                        break

            if self.first_target['set or rise'] == 'both':
                self.idx_second_target = self.idx_first_target
                self.second_target = self.first_target
                #print( self.first_target)
                break

        if self.idx_second_target is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no second target available')

    def table_priority_prio(self,day):

        self.priority=Table(names=('priority','target name','set or rise','alt set start','alt rise start','alt set end','alt rise end'),dtype=('f4','S11','S4','f4','f4','f4','f4'))
        try:
            self.observability_seclection(day) #observability selection
        except ValueError:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' This is a known error, please re-run')

        if (self.telescope == 'Io') or (self.telescope == 'Europa') or (self.telescope == 'Saint-Ex') or \
                (self.telescope == 'Artemis') or (self.telescope == 'Ganymede') or (self.telescope == 'Callisto') or\
                (self.telescope == 'TS_La_Silla') or (self.telescope == 'TN_Oukaimeden'):

            idx_on_going = np.where(((self.target_table_spc['nb_hours_surved'] > 0) & (self.target_table_spc['nb_hours_surved'] < 190)))
            idx_to_be_done = np.where((self.target_table_spc['nb_hours_surved'] == 0))
            idx_done = np.where((self.target_table_spc['nb_hours_surved'] > 200))
            idx_prog0 = np.where((self.target_table_spc['Program']==0))
            idx_prog1 = np.where((self.target_table_spc['Program']==1))
            idx_prog2 = np.where((self.target_table_spc['Program']==2))
            idx_prog3 = np.where((self.target_table_spc['Program'] == 3))
            self.priority['priority'][idx_prog0] *= 0.1
            self.priority['priority'][idx_prog1] *= 10 * self.target_table_spc['SNR_JWST_HZ_tr'][idx_prog1]**15
            self.priority['priority'][idx_prog2] *= 10 * self.target_table_spc['SNR_TESS_temp'][idx_prog2]** 0
            self.priority['priority'][idx_prog3] *= 10 * self.target_table_spc['SNR_Spec_temp'][idx_prog3] ** 0
            self.priority['priority'][idx_on_going] *= 10 ** (4 + 1 / (1 + 200 - self.target_table_spc['nb_hours_surved'][idx_on_going]))
            self.priority['priority'][idx_to_be_done] *= 10 ** (1 / (1 + 200 - self.target_table_spc['nb_hours_surved'][idx_to_be_done]))
            self.priority['priority'][idx_done] = -1

        set_targets_index = (self.priority['alt set start']>self.Altitude_constraint) & (self.priority['alt set end']>self.Altitude_constraint)
        self.priority['set or rise'][set_targets_index] = 'set'
        rise_targets_index = (self.priority['alt rise start']>20) & (self.priority['alt rise end']>self.Altitude_constraint)
        self.priority['set or rise'][rise_targets_index]='rise'
        both_targets_index=(self.priority['alt rise start']>self.Altitude_constraint) & (self.priority['alt set start']>self.Altitude_constraint) #& (self.priority['alt rise end']>20) & (self.priority['alt set end']>20)
        self.priority['set or rise'][both_targets_index] = 'both'
        self.priority['priority'][both_targets_index] = self.priority['priority'][both_targets_index]*10
        priority_non_observable_idx = (self.priority['set or rise']=='None')
        self.priority['priority'][priority_non_observable_idx] = 0.5

        if self.observatory.name == 'SSO':
            self.priority['priority'][self.idx_planned_SSO] = -1000
            self.priority['priority'][self.idx_planned_SNO] = -1000
            self.priority['priority'][self.idx_planned_TS] = -1000
            self.priority['priority'][self.idx_planned_TN] = -1000

        if (self.telescope == 'Artemis') or (self.telescope == 'Saint-Ex'):
            self.priority['priority'][self.idx_planned_SSO] = -1000
            self.priority['priority'][self.idx_planned_TS] = -1000
            self.priority['priority'][self.idx_planned_TN] = -1000

        if (self.telescope == 'TN_Oukaimeden'):
            self.priority['priority'][self.idx_planned_SSO] = -1000
            self.priority['priority'][self.idx_planned_SNO] = -1000
            self.priority['priority'][self.idx_planned_TS] = -1000

        if (self.telescope == 'TS_La_Silla'):
            self.priority['priority'][self.idx_planned_SSO] = -1000
            self.priority['priority'][self.idx_planned_SNO] = -1000
            self.priority['priority'][self.idx_planned_TN] = -1000

        self.no_obs_with_different_tel()
        try:
            read_exposure_time_table = pd.read_csv(path_spock + '/SPOCK_files/exposure_time_table.csv',sep=',')
        except FileNotFoundError:
            self.exposure_time_table(day)
            read_exposure_time_table = pd.read_csv(path_spock + '/SPOCK_files/exposure_time_table.csv', sep=',')
        if self.observatory.name == 'SSO':
            texp = read_exposure_time_table['SSO_texp']
            idx_texp_too_long = np.where((texp > 100))
            self.priority['priority'][idx_texp_too_long] = -1000
        if self.observatory.name == 'SNO':
            texp = read_exposure_time_table['SNO_texp']
            idx_texp_too_long =  np.where((texp > 90))
            self.priority['priority'][idx_texp_too_long] = -1000
        if self.observatory.name == 'Saint-Ex':
            texp = read_exposure_time_table['Saintex_texp']
            idx_texp_too_long =  np.where((texp > 110))
            self.priority['priority'][idx_texp_too_long] = -1000
        if self.observatory.name == 'TS_La_Silla':
            texp = read_exposure_time_table['TS_texp']
            idx_texp_too_long =  np.where((texp > 100))
            self.priority['priority'][idx_texp_too_long] = -1000
        if self.observatory.name == 'TN_Oukaimeden':
            texp = read_exposure_time_table['TN_texp']
            idx_texp_too_long =  np.where((texp > 100))
            self.priority['priority'][idx_texp_too_long] = -1000

        self.index_prio = np.argsort(self.priority['priority'])
        self.priority_ranked = self.priority[self.index_prio]
        # print(self.priority_ranked)

        return self.index_prio, self.priority, self.priority_ranked

    def observability_seclection(self, day):
        day_fmt = Time(day.iso, out_subfmt='date').iso
        if os.path.exists(path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name) + '_' + str(day_fmt) +  '_ndays_'  + str(self.date_range_in_days) + '_' + str(
                len(self.targets)) + '.csv'):
            name_file = path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name) + '_' + str(day_fmt) +  '_ndays_'  + str(self.date_range_in_days) + '_' + str(
                len(self.targets)) + '.csv'
            dataframe_ranking_months = pd.read_csv(name_file, delimiter=',')
            self.priority = Table.from_pandas(dataframe_ranking_months)
        else:
            start_night_start = self.observatory.twilight_evening_nautical(day, which='nearest') #* u.hour
            delta_midnight_start = np.linspace(0 , self.observatory.twilight_morning_nautical(day, which='next').jd - self.observatory.twilight_evening_nautical(day, which='nearest').jd , 100) * u.day
            frame_start = AltAz(obstime = start_night_start + delta_midnight_start, location=self.observatory.location)
            #target_alt_start = [(target.coord.transform_to(frame_start).alt) for target in self.targets]
            start_night_end = self.observatory.twilight_evening_nautical(day + self.date_range_in_days, which='nearest') #* u.hour
            delta_midnight_end = np.linspace(0 , self.observatory.twilight_morning_nautical(day + self.date_range_in_days, which='next').jd - self.observatory.twilight_evening_nautical(day + self.date_range_in_days, which='nearest').jd , 100) * u.day
            frame_end = AltAz(obstime = start_night_end + delta_midnight_end, location=self.observatory.location)
            target_alt = [[target.coord.transform_to(frame_start).alt,target.coord.transform_to(frame_end).alt] for target in self.targets]
            target_alt_start = np.asarray(target_alt)[:,0,:]
            target_alt_end = np.asarray(target_alt)[:,1,:]
            #
            max_target_alt = list(map(max, target_alt_start))
            alt_set_start = list(map(first_elem_list, target_alt_start[:]))
            alt_rise_start = list(map(last_elem_list, target_alt_start[:]))
            alt_set_end = list(map(first_elem_list, target_alt_end[:]))
            alt_rise_end = list(map(last_elem_list, target_alt_end[:]))
            priority = [-0.5] * len(self.targets)
            set_or_rise = ['None'] * len(self.targets)
            df = pd.DataFrame({'priority':priority,'set or rise':set_or_rise,'alt set start':alt_set_start, 'alt rise start':alt_rise_start, 'alt set end':alt_set_end, 'alt rise end':alt_rise_end,'max_alt':max_target_alt,'Sp_ID':self.target_table_spc['Sp_ID']})
            self.priority = Table.from_pandas(df)
            month_opt = Table([[], [], [], [], []], names=['months', 'months_2nd', 'months_3rd', 'months_4th', 'months_5th'])
            [month_opt.add_row(month_option(target, self.reverse_df1)) for target in self.target_table_spc['Sp_ID']]
            idx_1rst_opt_monthobs = np.where((month_opt['months']==self.months_obs))
            idx_2nd_opt_monthobs = np.where((month_opt['months_2nd'] == self.months_obs))
            idx_3rd_opt_monthobs = np.where((month_opt['months_3rd'] == self.months_obs))
            idx_4th_opt_monthobs = np.where((month_opt['months_4th'] == self.months_obs))
            idx_5th_opt_monthobs = np.where((month_opt['months_5th'] == self.months_obs))
            self.priority['priority'][idx_1rst_opt_monthobs] = (self.priority['max_alt'][idx_1rst_opt_monthobs] - 30 ) *10**4
            self.priority['priority'][idx_2nd_opt_monthobs] = (self.priority['max_alt'][idx_2nd_opt_monthobs] - 30 ) *10**3
            self.priority['priority'][idx_3rd_opt_monthobs] = (self.priority['max_alt'][idx_3rd_opt_monthobs] - 30 ) *10**2
            self.priority['priority'][idx_4th_opt_monthobs] = (self.priority['max_alt'][idx_4th_opt_monthobs] - 30 ) *10**1
            self.priority['priority'][idx_5th_opt_monthobs] = (self.priority['max_alt'][idx_5th_opt_monthobs] - 30 ) *10**0

            dataframe_priority = self.priority.to_pandas()
            dataframe_priority.to_csv(path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name) + '_' + str(day_fmt) +  '_ndays_'  + str(self.date_range_in_days) + '_' + str(len(self.targets)) + '.csv', sep=',', index=False)

    def shift_hours_observation(self,idx_target):

        date_format = "%Y-%m-%d %H:%M:%S.%f"

        if self.priority['set or rise'][idx_target] == 'set':

            set_time_begin = datetime.strptime(self.observatory.target_set_time(self.date_range[0],self.targets[idx_target],which='next',horizon=24*u.deg).iso, date_format)
            set_time_end = datetime.strptime(self.observatory.target_set_time(self.date_range[1],self.targets[idx_target],which='next',horizon=24*u.deg).iso, date_format)
            shift_hours_observation = (set_time_begin.hour + set_time_begin.minute/60 - set_time_end.hour - set_time_end.minute/60 )

        if self.priority['set or rise'][idx_target] == 'rise':

            rise_time_begin = datetime.strptime(self.observatory.target_rise_time(self.date_range[0],self.targets[idx_target],which='next',horizon=24*u.deg).iso, date_format)
            rise_time_end = datetime.strptime(self.observatory.target_rise_time(self.date_range[1],self.targets[idx_target],which='next',horizon=24*u.deg).iso, date_format)
            shift_hours_observation = (rise_time_end.hour + rise_time_end.minute/60 - rise_time_begin.hour - rise_time_begin.minute/60)

        if self.priority['set or rise'][idx_target] == 'both':
            shift_hours_observation = 0

        return shift_hours_observation #hours

    def schedule_blocks(self,day):
        """
            schedule the lock thanks to astroplan tools

        Parameters
        ----------
            idx_first_target: int, index of the first target
            idx_second_target: int, index of the second target
        Returns
        -------
            SS1_night_blocks: astropy.table with name, start time, end time, duration ,
            coordinates (RE and DEC) and configuration (filter and exposure time) for each target of the night

        """

        dt_1day=Time('2018-01-02 00:00:00',scale='tcg')-Time('2018-01-01 00:00:00',scale='tcg')
        delta_day = (day - self.date_range[0]).value

        if self.idx_second_target is not None:
            shift = max(self.shift_hours_observation(self.idx_first_target),self.shift_hours_observation(self.idx_second_target)) /24 #days
        else:
            shift = self.shift_hours_observation(self.idx_first_target) / 24 #days

        dur_obs_both_target = (self.night_duration(day).value)* u.day
        dur_obs_set_target = (self.night_duration(day).value/2 - shift/self.date_range_in_days)  * u.day #(self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
        dur_obs_rise_target = (self.night_duration(day).value/2 + shift/self.date_range_in_days) * u.day  #(self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)

        constraints_set_target = self.constraints + [TimeConstraint((Time(self.observatory.twilight_evening_nautical(self.date_range[0]+ dt_1day * delta_day,which='next'))),
                                                                    (Time(self.observatory.twilight_evening_nautical(self.date_range[0]+ dt_1day * delta_day,which='next') + dur_obs_set_target)))]

        constraints_rise_target = self.constraints + [TimeConstraint((Time(self.observatory.twilight_evening_nautical(self.date_range[0] + dt_1day * delta_day,which='next') + dur_obs_set_target)),
                                                                     (Time(self.observatory.twilight_morning_nautical(self.date_range[0] + dt_1day * (delta_day+1),which='nearest') )))]

        constraints_all = self.constraints + [TimeConstraint((Time(self.observatory.twilight_evening_nautical(self.date_range[0] + dt_1day * delta_day,which='next'))),
                                                             (Time(self.observatory.twilight_morning_nautical(self.date_range[0] + dt_1day * (delta_day+1),which='nearest'))))]

        blocks=[]
        if self.target_table_spc['texp_spc'][self.idx_first_target] == 0:
            self.target_table_spc['texp_spc'][self.idx_first_target]= self.exposure_time(day=None,i=self.idx_first_target)
        if self.telescope == 'Saint-Ex' or self.telescope == 'Artemis':  # different default so bbetter to recalculate
            self.target_table_spc['texp_spc'][self.idx_first_target] = self.exposure_time(day=None, i=self.idx_first_target)
        if (self.idx_second_target != None) and self.target_table_spc['texp_spc'][self.idx_second_target] == 0:
            self.target_table_spc['texp_spc'][self.idx_second_target] = self.exposure_time(day=None,i=self.idx_second_target)
        if self.telescope == 'Saint-Ex' or self.telescope == 'Artemis':  # different default so bbetter to recalculate
            self.target_table_spc['texp_spc'][self.idx_second_target] = self.exposure_time(day=None, i=self.idx_second_target)

        if self.first_target['set or rise'] == 'set':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'set\'')
            a = ObservingBlock(self.targets[self.idx_first_target],dur_obs_set_target,-1,constraints= constraints_set_target,configuration={'filt=' + str(self.target_table_spc['Filter_spc'][self.idx_first_target]),'texp=' + str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)
            b = ObservingBlock(self.targets[self.idx_second_target],dur_obs_rise_target,-1,constraints=constraints_rise_target,configuration={'filt=' + str(self.target_table_spc['Filter_spc'][self.idx_second_target]),'texp=' + str(self.target_table_spc['texp_spc'][self.idx_second_target])})
            blocks.append(b)

        if self.first_target['set or rise'] == 'both':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'both\'')
            a = ObservingBlock(self.targets[self.idx_first_target],dur_obs_both_target,-1,constraints= constraints_all,configuration={'filt=' + str(self.target_table_spc['Filter_spc'][self.idx_first_target]),'texp=' + str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)

        if self.first_target['set or rise'] == 'rise':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'rise\'')
            b = ObservingBlock(self.targets[self.idx_second_target],dur_obs_set_target,-1,constraints=constraints_set_target,configuration={'filt=' + str(self.target_table_spc['Filter_spc'][self.idx_second_target]),'texp=' + str(self.target_table_spc['texp_spc'][self.idx_second_target])})
            blocks.append(b)
            a = ObservingBlock(self.targets[self.idx_first_target],dur_obs_rise_target,-1,constraints=constraints_rise_target,configuration={'filt=' + str(self.target_table_spc['Filter_spc'][self.idx_first_target]),'texp=' + str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)

        transitioner = Transitioner(slew_rate= 11*u.deg/u.second)
        seq_schedule_SS1 = Schedule(self.date_range[0] + dt_1day * delta_day ,self.date_range[0] + dt_1day * (delta_day + 1))
        sequen_scheduler_SS1 = SPECULOOSScheduler(constraints=constraints_all, observer=self.observatory,transitioner=transitioner)
        sequen_scheduler_SS1(blocks,seq_schedule_SS1)

        SS1_night_blocks = seq_schedule_SS1.to_table()
        name_all = SS1_night_blocks['target']
        name=[]
        for i,nam in enumerate(name_all):
            name.append(nam)
        return SS1_night_blocks

    def make_night_block(self,day):

        day_fmt = Time(day.iso, out_subfmt='date').iso

        self.night_block.add_index('target')

        try:
            index_to_delete = self.night_block.loc['TransitionBlock'].index
            self.night_block.remove_row(index_to_delete)
        except KeyError:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transition block')

        panda_table = self.night_block.to_pandas()
        panda_table.to_csv(os.path.join(path_spock + '/night_blocks_propositions/' +'night_blocks_' + self.telescope + '_' +  str(day_fmt) + '.txt'),sep=' ')

    def is_constraint_hours(self,idx_target):
        """
            Check if number of hours is ok

        Parameters
        ----------
            idx_target: int, index of the target you want to check

        Returns
        -------
            is_hours_constraint_met_target: boolean, say the hour constraint is ok or not

        """
        nb_hours_observed = self.target_table_spc['nb_hours_surved']
        is_hours_constraint_met_target = True
        a=(1-self.target_table_spc['nb_hours_threshold'][idx_target]/(self.target_table_spc['nb_hours_threshold'][idx_target]+10))
        if a < 1E-5:
            is_hours_constraint_met_target = False
        return is_hours_constraint_met_target

    def night_duration(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        dt_1day=Time('2018-01-02 00:00:00',scale='tt')-Time('2018-01-01 00:00:00',scale='tt')
        dura = Time(Time(self.observatory.twilight_morning_nautical(day + dt_1day ,which='nearest')).jd - \
               Time(self.observatory.twilight_evening_nautical(day ,which='next')).jd,format='jd')
        return dura

    def info_obs_possible(self,day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        duration_obs_possible = self.is_moon_and_visibility_constraint(day)['fraction of time observable'] # * Time(self.night_duration(day)).value  #np.subtract(np.asarray(self.set_time_targets(day).value) , np.asarray(self.rise_time_targets(day).value))
        duration_obs_possible[duration_obs_possible <0]= 0
        return duration_obs_possible

    def rise_time_targets(self,day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        rise_time = self.observatory.target_rise_time(day, self.targets, which='next', horizon=24 * u.deg)
        return rise_time

    def set_time_targets(self,day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')
        set_time = self.observatory.target_set_time(day+1*dt_1day,self.targets, which='nearest', horizon=24 * u.deg)
        return set_time

    def idx_is_julien_criterion(self,day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        nb_hour_wanted = 6 #hours
        idx_is_julien_criterion = (self.info_obs_possible(day) > nb_hour_wanted )
        return idx_is_julien_criterion

    def is_moon_and_visibility_constraint(self, day):
        """
            Check if number of moon not too close and target is visible

        Parameters
        ----------
            idx_target: int, index of the target you want to check
        Returns
        -------
            is_hours_constraint_met_target: boolean, say the hour constraint is ok or not

        """
        constraints = self.constraints + [MoonSeparationConstraint(min=35*u.deg)]

        dt_1day=Time('2018-01-02 00:00:00',scale='utc')-Time('2018-01-01 00:00:00',scale='utc')

        self.observability_table_day = observability_table(self.constraints, self.observatory, self.targets,
         time_range = Time([Time(self.observatory.twilight_evening_nautical(Time(day,format='jd'), which='next')).iso,
                            Time(self.observatory.twilight_morning_nautical(Time(day + dt_1day,format='jd'),
                                                                            which='nearest')).iso])) # + nd/2

        self.observability_table_day['fraction of time observable'] =  self.observability_table_day['fraction of time observable'] * \
                                                                       Time(self.night_duration(day)).value * 24

        is_visible_mid_night = is_observable(constraints, self.observatory, self.targets, \
            times= Time(self.observatory.twilight_evening_nautical(Time(day),which='next') +
                        self.night_duration(day).value/2))

        idx_not_visible_mid_night = np.where((is_visible_mid_night == False))

        self.observability_table_day['ever observable'][idx_not_visible_mid_night[0]] = False

        return self.observability_table_day

    def is_constraints_met_first_target(self,t):
        """
            Useful when the moon constrain (< 30°) or the hours constraint (nb_hours_observed>nb_hours_threshold)
            are not fullfilled anymore and the first target needs to be changed

        Parameters
        ----------
            t: int, days in the month
            idx_first_target: int, index associated to the first target (with regards to the targets list)
            idx_set_targets_sorted: list of int, index of all the set target ranked y priority
            idx_rise_targets_sorted: list of int, index of all the rise target ranked y priority
            index_prio: list of int, index of all the rise target ranked y priority

        Returns
        -------
            is_moon_constraint_met_first_target: boolean, say if moon constraint is fullfilled on first target
            hours_constraint: boolean, say if hour constraint is fullfilled on first target
            idx_first_target: int, index for the first target (if constraints where already fullfilled the Value
            is the same as input, else it is a new value for a new target)

        """
        dt_1day=Time('2018-01-02 00:00:00',scale='tcg')-Time('2018-01-01 00:00:00',scale='tcg')

        is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
        hours_constraint_first = self.is_constraint_hours(self.idx_first_target)

        idx_safe=1
        idx_safe_1rst_set=1
        idx_safe_1rst_rise=1
        moon_idx_set_target=0
        moon_idx_rise_target=0

        while not (is_moon_constraint_met_first_target & hours_constraint_first):

            beFore_change_first_target=self.priority[self.idx_first_target]
            beFore_change_idx_first_target=self.idx_first_target

            if beFore_change_first_target['set or rise'] == 'set':
                moon_idx_set_target += 1
                if moon_idx_set_target >= len(self.idx_set_targets_sorted):
                    idx_safe_1rst_set += 1
                    self.idx_first_target = self.index_prio[-idx_safe_1rst_set]
                    self.first_target = self.priority[self.idx_first_target]
                    is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table ['ever observable'][self.idx_first_target]
                    hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
                else:
                    self.idx_first_target=self.idx_set_targets_sorted[-moon_idx_set_target]
                    self.first_target=self.priority[self.idx_first_target]
                    if self.priority['priority'][self.idx_first_target]!=float('-inf'):
                        is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
                        hours_constraint_first=self.is_constraint_hours(self.idx_first_target)
                    else:
                        is_moon_constraint_met_first_target = False
                        hours_constraint_first = False

            if beFore_change_first_target['set or rise'] == 'rise':
                moon_idx_rise_target += 1
                if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
                    idx_safe_1rst_rise += 1
                    self.idx_first_target = self.index_prio[-idx_safe_1rst_rise]
                    self.first_target = self.priority[self.idx_first_target]
                    is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
                    hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
                else:
                    self.idx_first_target=self.idx_rise_targets_sorted[-moon_idx_rise_target]
                    self.first_target = self.priority[self.idx_first_target]
                    if self.priority['priority'][self.idx_first_target] != float('-inf'):
                        is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
                        hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
                    else:
                        is_moon_constraint_met_first_target = False
                        hours_constraint_first = False


            if (beFore_change_first_target['set or rise'] != 'rise') and (beFore_change_first_target['set or rise'] != 'set'):
                idx_safe+=1
                self.idx_first_target=self.index_prio[-idx_safe]
                self.first_target=self.priority[self.idx_first_target]
                is_moon_constraint_met_first_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
                hours_constraint_first=self.is_constraint_hours(self.idx_first_target)

        if is_moon_constraint_met_first_target and hours_constraint_first:
            is_constraints_met_first_target=True
        else:
            is_constraints_met_first_target=False
        return is_constraints_met_first_target

    def is_constraints_met_second_target(self,t):
        """
            Useful when the moon constrain (< 30°) or the hours constraint (nb_hours_observed>nb_hours_threshold)
            are not fullfilled anymore and the second target needs to be changed

        Parameters
        ----------
            t: int, days in the month
            idx_first_target: int, index associated to the first target (with regards to the targets list)
            idx_second_target: int, index associated to the second target (with regards to the targets list)
            idx_set_targets_sorted: list of int, index of all the set target ranked y priority
            idx_rise_targets_sorted: list of int, index of all the rise target ranked y priority
            index_prio: list of int, index of all the rise target ranked y priority

        Returns
        -------
            is_moon_constraint_met_second_target: boolean, say if moon constraint is fullfilled on second target
            hours_constraint: boolean, say if hour constraint is fullfilled on second target
            idx_second_target: int, index for the second target (if constraints where already fullfilled the Value
            is the same as input, else it is a new value for a new target)

        """
        dt_1day=Time('2018-01-02 00:00:00',scale='tcg')-Time('2018-01-01 00:00:00',scale='tcg')

        is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
        hours_constraint_second=self.is_constraint_hours(self.idx_second_target)

        idx_safe=1
        idx_safe_2nd_set=1
        idx_safe_2nd_rise=1
        moon_idx_set_target=0
        moon_idx_rise_target=0

        if self.idx_first_target == self.idx_second_target:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' 2nd = 1ere ')
        if self.idx_second_target is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' No second target for that night')
        else:
            while not (is_moon_constraint_met_second_target & hours_constraint_second):

                beFore_change_second_target = self.priority[self.idx_second_target]

                if beFore_change_second_target['set or rise'] == 'set':
                    moon_idx_set_target += 1
                    if moon_idx_set_target >= len(self.idx_set_targets_sorted):
                        idx_safe_2nd_set += 1
                        self.idx_second_target = self.index_prio[-idx_safe_2nd_set]
                        self.second_target = self.priority[self.idx_second_target]
                        is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                        hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
                    else:
                        self.idx_second_target=self.idx_set_targets_sorted[-(moon_idx_set_target)]
                        self.second_target=self.priority[self.idx_second_target]
                        if self.priority['priority'][self.idx_second_target] != float('-inf'):
                            is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                            hours_constraint_second=self.is_constraint_hours(self.idx_second_target)
                        else:
                            is_moon_constraint_met_second_target=False
                            hours_constraint_second = False

                if beFore_change_second_target['set or rise'] == 'rise':
                    moon_idx_rise_target += 1
                    if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
                        idx_safe_2nd_rise += 1
                        self.idx_second_target = self.index_prio[-idx_safe_2nd_rise]
                        self.second_target = self.priority[self.idx_second_target]
                        is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                        hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
                    else:
                        self.idx_second_target = self.idx_rise_targets_sorted[-(moon_idx_rise_target)]
                        self.second_target = self.priority[self.idx_second_target]
                        if self.priority['priority'][self.idx_second_target] != float('-inf'):
                            is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                            hours_constraint_second=self.is_constraint_hours(self.idx_second_target)
                        else:
                            is_moon_constraint_met_second_target=False
                            hours_constraint_second = False

                if (beFore_change_second_target['set or rise'] != 'rise') and (beFore_change_second_target['set or rise'] != 'set'):
                    if self.first_target['set or rise'] == 'rise':
                        moon_idx_set_target+=1
                        if moon_idx_set_target >= len(self.idx_set_targets_sorted):
                            idx_safe_2nd_set += 1
                            self.idx_second_target = self.index_prio[-idx_safe_2nd_set]
                            self.second_target = self.priority[self.idx_second_target]
                            is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                            hours_constraint_second=self.is_constraint_hours(self.idx_second_target)
                        else:
                            self.idx_second_target = self.idx_set_targets_sorted[-(moon_idx_set_target)]
                            self.second_target = self.priority[self.idx_second_target]
                            if self.priority['priority'][self.idx_second_target] != float('-inf'):
                                is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                                hours_constraint_second=self.is_constraint_hours(self.idx_second_target)
                            else:
                                is_moon_constraint_met_second_target=False
                                hours_constraint_second = False

                    if self.first_target['set or rise'] == 'set':
                        moon_idx_rise_target+=1
                        if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
                            idx_safe_2nd_rise += 1
                            self.idx_second_target = self.index_prio[-idx_safe_2nd_rise]
                            self.second_target=self.priority[self.idx_second_target]
                            is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                            hours_constraint_second=self.is_constraint_hours(self.idx_second_target)

                        else:
                            self.idx_second_target = self.idx_rise_targets_sorted[-(moon_idx_rise_target)]
                            self.second_target = self.priority[self.idx_second_target]
                            if self.priority['priority'][self.idx_second_target] != float('-inf'):
                                is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                                hours_constraint_second=self.is_constraint_hours(self.idx_second_target)
                            else:
                                is_moon_constraint_met_second_target=False
                                hours_constraint_second = False

                    if self.first_target['set or rise'] == 'both':
                        self.idx_second_target = self.idx_first_target
                        self.second_target = self.first_target
                        #is_moon_constraint_met_second_target = self.is_moon_and_visibility_constraint(t)
                        is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
                        hours_constraint_second = self.is_constraint_hours(self.idx_second_target)

        if is_moon_constraint_met_second_target and hours_constraint_second:
            is_constraints_met_second_target = True
        else:
            is_constraints_met_second_target = False
        return is_constraints_met_second_target

    def update_hours_observed_first(self,day):
        """
            update number of hours observed for the corresponding first target

        Parameters
        ----------
            t: int, day of the month
            idx_first_target: int, index of the first target
        Returns
        -------
            self.nb_hours_observed
            self.nb_hours
            self.nb_hours[idx_first_target]

        """
        if self.idx_second_target is not None:
            shift = max(self.shift_hours_observation(self.idx_first_target),self.shift_hours_observation(self.idx_second_target)) /24 #days
        else:
            shift = self.shift_hours_observation(self.idx_first_target) /24 #days

        dur_obs_both_target = (self.night_duration(day).value) * u.day
        dur_obs_set_target = (self.night_duration(day).value/2 - shift/self.date_range_in_days) * u.day #(self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
        dur_obs_rise_target = (self.night_duration(day).value/2 + shift/self.date_range_in_days) * u.day#(self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)

        if self.first_target['set or rise'] == 'set':
            nb_hours__1rst_old = self.nb_hours[1,self.idx_first_target] #hours
            a =  dur_obs_set_target.value #in days

        if self.first_target['set or rise'] == 'both':
            nb_hours__1rst_old = self.nb_hours[1,self.idx_first_target] #hours
            a  = dur_obs_both_target.value #in days

        if self.first_target['set or rise'] == 'rise':
            nb_hours__1rst_old = self.nb_hours[1,self.idx_first_target] #hours
            a = dur_obs_rise_target.value #in days

        self.nb_hours[0,self.idx_first_target] = nb_hours__1rst_old
        self.nb_hours[1,self.idx_first_target] = nb_hours__1rst_old +  a *24
        self.target_table_spc['nb_hours_surved'][self.idx_first_target] = nb_hours__1rst_old +  a *24

    def update_hours_observed_second(self,day):
        """
            update number of hours observed for the corresponding second target

        Parameters
        ----------
            t: int, day of the month
            idx_second_target: int, index of the second target
        Returns
        -------
            self.nb_hours_observed
            self.nb_hours
            self.nb_hours[idx_second_target]

        """
        if self.idx_second_target is not None:
            shift = max(self.shift_hours_observation(self.idx_first_target),self.shift_hours_observation(self.idx_second_target)) /24 #days
        else:
            shift = self.shift_hours_observation(self.idx_first_target) /24 #days
        dur_obs_both_target = (self.night_duration(day).value ) * u.day
        dur_obs_set_target = (self.night_duration(day).value/2  - shift/self.date_range_in_days)  * u.day #(self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
        dur_obs_rise_target = (self.night_duration(day).value/2  + shift/self.date_range_in_days) * u.day #(self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)


        if self.second_target['set or rise'] == 'rise':
            nb_hours__2nd_old = self.nb_hours[1,self.idx_second_target] #hours
            b = dur_obs_rise_target.value #in days
        if self.second_target['set or rise'] == 'set':
            nb_hours__2nd_old = self.nb_hours[1,self.idx_second_target] #hours
            b =  dur_obs_set_target.value #in days
        if self.second_target['set or rise'] == 'both':
            nb_hours__2nd_old = self.nb_hours[1,self.idx_second_target] #hours
            b = dur_obs_both_target.value #in days

        self.nb_hours[0, self.idx_second_target] = nb_hours__2nd_old
        self.nb_hours[1, self.idx_second_target] = nb_hours__2nd_old + b *24
        self.target_table_spc['nb_hours_surved'][self.idx_second_target] =  nb_hours__2nd_old + b *24

    def make_plan_file(self,day):
        name_all = self.night_block['target']
        Start_all = self.night_block['start time (UTC)']
        Finish_all = self.night_block['end time (UTC)']
        Duration_all = self.night_block['duration (minutes)']
        self.update_hours_observed_first(day)
        if self.idx_second_target is not None:
            self.update_hours_observed_second(day)
        file_txt = 'plan.txt'
        with open(file_txt, 'a') as file_plan:
            for i, nam in enumerate(name_all):
                if i == 1:
                    if nam == name_all[i - 1]:
                        Start_all[i] = Start_all[i - 1]
                        Duration_all[i] = (Time(Finish_all[i]) - Time(Start_all[i])).value * 24 * 60
                        self.night_block.remove_row(0)

                idx_target = np.where((nam == self.target_table_spc['Sp_ID']))[0]
                if nam != 'TransitionBlock':
                    file_plan.write(
                        nam + ' ' + '\"' + Start_all[i] + '\"' + ' ' + '\"' + Finish_all[i] + '\"' + ' ' + self.telescope + ' ' + \
                        str(np.round(self.nb_hours[0,idx_target[0]], 3)) + '/' + str(
                            self.target_table_spc['nb_hours_threshold'][idx_target[0]]) + ' ' + \
                        str(np.round(self.nb_hours[1,idx_target[0]], 3)) + '/' + str(
                            self.target_table_spc['nb_hours_threshold'][idx_target[0]]) + '\n')

    def reference_table(self):

        a = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        b = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        c = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        d = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        idx_true = [0] * len(self.date_ranges_day_by_day)
        is_julien_criterion = [False] * len(self.date_ranges_day_by_day)

        for i, day in enumerate(self.date_ranges_day_by_day):
            a[i] = self.is_moon_and_visibility_constraint(day)['ever observable']
            b[i] = self.is_moon_and_visibility_constraint(day)['fraction of time observable']
            idx_true = len(np.where(a[i])[0])
            is_julien_criterion[i] = np.sum(self.idx_is_julien_criterion(day))
            c[i] = self.rise_time_targets(day)
            d[i] = self.set_time_targets(day)

        E = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']), 5)).astype("object")
        E[:, :, 0] = self.target_table_spc['Sp_ID']
        E[:, :, 1] = a
        E[:, :, 2] = b
        E[:, :, 3] = c
        E[:, :, 4] = d

        np.save('array' + str(self.observatory.name) + '_6hr'+ 'prio_50_no_M6' + '.npy', E, allow_pickle=True)

        df = pd.DataFrame({'day': Time(self.date_ranges_day_by_day), 'nb_observable_target': idx_true,
                           'julien_criterion': is_julien_criterion})

        df.to_csv('nb_observable_target_prio_50_' + str(self.observatory.name) + '_6hr' + 'prio_50_no_M6' + '.csv', sep=',', index=False)

    def exposure_time_for_table(self, obs,day,i): #juste for exposure time table
        """ calcul of exposure time to store in a table

        Parameters
        ----------
        obs : str
            name of observatory
        day : date
            day for moon phse calcultation (fmt 'yyyy-mm-dd')
        i : int
            target index in target list

        Returns
        -------
        float
            exposure time

        """

        if day is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
        #moon_phase = round(moon_illumination(Time(day.iso, out_subfmt='date')), 2)
        if round(self.target_table_spc['SpT'][i]) <= 9:
            if round(self.target_table_spc['SpT'][i]) <= 0:
                spt_type = 'M0'
            else:
                spt_type = 'M' + str(round(self.target_table_spc['SpT'][i]))
                if spt_type == 'M3':
                    spt_type = 'M2'
                if spt_type == 'M1':
                    spt_type = 'M0'
        elif (round(self.target_table_spc['SpT'][i]) == 12) or (round(self.target_table_spc['SpT'][i]) == 15) or (
                round(self.target_table_spc['SpT'][i]) == 18):
            spt_type = 'M' + str(round(self.target_table_spc['SpT'][i]) - 10)
        elif round(self.target_table_spc['SpT'][i]) == 10:
            spt_type = 'M9'
        elif round(self.target_table_spc['SpT'][i]) == 11:
            spt_type = 'L2'
        elif round(self.target_table_spc['SpT'][i]) == 13:
            spt_type = 'L2'
        elif round(self.target_table_spc['SpT'][i]) == 14:
            spt_type = 'L5'
        elif round(self.target_table_spc['SpT'][i]) > 14:
            spt_type = 'L8'
        filt_ = str(self.target_table_spc['Filter_spc'][i])
        if (filt_ == 'z\'') or (filt_ == 'r\'') or (filt_ == 'i\'') or (filt_ == 'g\''):
            filt_ = filt_.replace('\'', '')
        filters = ['I+z']
        filt_idx = 0
        filt_ = filters[filt_idx]
        if obs =='Saint-Ex':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1,
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=3.48, temp_ccd=-70,
                         observatory_altitude=2780))
            texp = a.exp_time_calculator(ADUpeak=30000)[0]
        elif obs == 'TS' or obs =='TN':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1,\
                        moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=50000)[0]
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')
        elif obs =='SNO':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1,\
                        moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=45000)[0]
        elif obs == 'SSO':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=45000)[0]

        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' texp is : ',texp)
        return texp

    def exposure_time(self, day, i):
        """ calculation of the exposure time for a given target

        Parameters
        ----------
        day : date
            format 'yyyy-mm-dd'
        i : int
            index of target in target_list

        Returns
        -------
        float
            exposure time

        """
        if day is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
        # moon_phase = round(moon_illumination(Time(day.iso, out_subfmt='date')), 2)
        if round(self.target_table_spc['SpT'][i]) <= 9:
            spt_type = 'M' + str(round(self.target_table_spc['SpT'][i]))
            if spt_type == 'M3':
                spt_type = 'M2'
        elif (round(self.target_table_spc['SpT'][i]) == 12) or (round(self.target_table_spc['SpT'][i]) == 15) or (
                int(self.target_table_spc['SpT'][i]) == 18):
            spt_type = 'M' + str(round(self.target_table_spc['SpT'][i]) - 10)
        elif round(self.target_table_spc['SpT'][i]) == 10:
            spt_type = 'M9'
        elif round(self.target_table_spc['SpT'][i]) == 11:
            spt_type = 'L2'
        elif round(self.target_table_spc['SpT'][i]) == 13:
            spt_type = 'L2'
        elif round(self.target_table_spc['SpT'][i]) == 14:
            spt_type = 'L5'
        elif round(self.target_table_spc['SpT'][i]) > 14:
            spt_type = 'L8'
        filt_ = str(self.target_table_spc['Filter_spc'][i])
        if (filt_ == 'z\'') or (filt_ == 'r\'') or (filt_ == 'i\'') or (filt_ == 'g\''):
            filt_ = filt_.replace('\'', '')
        filters = ['I+z', 'z', 'i', 'r']
        filt_idx = 0
        filt_ = filters[filt_idx]
        if self.telescope == 'Saint-Ex':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1,
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=3.48, temp_ccd=-70,
                         observatory_altitude=2780))
            texp = a.exp_time_calculator(ADUpeak=30000)[0]
        elif self.telescope == 'TS_La_Silla' or self.telescope == 'TN_Oukaimeden':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=50000)[0]
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')
        elif self.telescope == 'Artemis':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=45000)[0]
        elif self.telescope == 'Io' or self.telescope == 'Europa' \
                or self.telescope == 'Ganymede' or self.telescope == 'Callisto':
            a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                         moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
            texp = a.exp_time_calculator(ADUpeak=45000)[0]

        while texp < 10:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Change filter to avoid saturation!!')
            filt_idx += 1
            if (filt_idx >= 3):
                print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' You have to defocus in we want to observe this target')
                texp = 10.0001
            filt_ = filters[filt_idx]
            if self.telescope == 'Saint-Ex':
                self.target_table_spc['Filter_spc'][i] = filt_
                a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1,
                             moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=3.48,temp_ccd=-70,
                             observatory_altitude=2780))
                texp = a.exp_time_calculator(ADUpeak=30000)[0]
            elif self.telescope == 'Artemis':
                self.target_table_spc['Filter_spc'][i] = filt_
                a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                             moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
                texp = a.exp_time_calculator(ADUpeak=45000)[0]
            elif self.telescope == 'TS_La_Silla' or self.telescope == 'TN_Oukaimeden':
                self.target_table_spc['Filter_spc'][i] = filt_
                a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                             moonphase=0.6, irtf=0.8, num_tel=1, seeing=1.05, gain=1.1))
                texp = a.exp_time_calculator(ADUpeak=50000)[0]
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')
            elif self.telescope == 'Io' or self.telescope == 'Europa' or self.telescope == 'Ganymede' or self.telescope == 'Callisto':
                self.target_table_spc['Filter_spc'][i] = filt_
                a = (ETC.etc(mag_val=self.target_table_spc['J'][i], mag_band='J', spt=spt_type, filt=filt_, airmass=1.1, \
                             moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
                texp = a.exp_time_calculator(ADUpeak=45000)[0]

        return texp

    def no_obs_with_different_tel(self):
        """ function to avoid observations of a similar target,
        expect for SNO and SAINT-EX were similar target observations are encouraged.

        Returns
        -------
        self.priority
            fill with 0 priority of targets already scheduled

        """
        idx_observed_SSO = self.idx_SSO_observed_targets()
        idx_observed_SNO = self.idx_SNO_observed_targets()
        idx_observed_SaintEx = self.idx_SaintEx_observed_targets()
        idx_observed_trappist = self.idx_trappist_observed_targets()
        if (self.telescope == 'TS_La_Silla') :
            #self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            self.priority['priority'][idx_observed_SaintEx] = 0
        elif (self.telescope == 'TN_Oukaimeden') :
            self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            self.priority['priority'][idx_observed_SaintEx] = 0
        elif (self.telescope == 'Artemis') or (self.telescope == 'Saint-Ex'):
            #self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_SaintEx] *= 2
            self.priority['priority'][idx_observed_SNO] *= 2
            self.priority['priority'][idx_observed_trappist] = 0
        elif (self.telescope == 'Io') or (self.telescope == 'Europa') or \
                (self.telescope == 'Ganymede') or (self.telescope == 'Callisto'):
            #self.priority['priority'][idx_observed_trappist] = 0
            self.priority['priority'][idx_observed_SaintEx] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            for i in range(len(idx_observed_SSO)):
                if self.target_table_spc['telescope'][idx_observed_SSO][i] != '':
                    if self.telescope not in self.target_table_spc['telescope'][idx_observed_SSO][i]:
                        self.priority['priority'][idx_observed_SSO[i]] = 0
        print()

def read_night_block(telescope, day):
    """[summary]

    Parameters
    ----------
    telescope : [type]
        [description]
    day : [type]
        [description]

    Returns
    -------
    [type]
        [description]
    """

    day_fmt = Time(day, scale='utc', out_subfmt='date').tt.datetime.strftime("%Y-%m-%d")
    path_local = path_spock + '/DATABASE/' + telescope + '/Archive_night_blocks/night_blocks_' + \
                 telescope + '_' + day_fmt + '.txt'

    if os.path.exists(path_local):
        day_fmt = Time(day, scale='utc', out_subfmt='date').tt.datetime.strftime("%Y-%m-%d")
        scheduler_table = Table.read(
            path_spock + '/DATABASE/' + str(telescope) + '/Archive_night_blocks' + '/night_blocks_' + str(telescope) + '_' + str(
                day_fmt) + '.txt',
            format='ascii')
    else:
        nightb_url = "http://www.mrao.cam.ac.uk/SPECULOOS/" + telescope + '/schedule/Archive_night_blocks/night_blocks_' + \
                     telescope + '_' + day_fmt + '.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            open(path_local, 'wb').write(nightb.content)
            scheduler_table = pd.read_csv(path_local, delimiter=' ',
                                          skipinitialspace=True, error_bad_lines=False)

    return scheduler_table


def make_docx_schedule(observatory,telescope, date_range, name_operator):
    if not os.path.exists(path_spock + '/TRAPPIST_schedules_docx'):
        os.makedirs(path_spock + '/TRAPPIST_schedules_docx')
    df_speculoos = pd.read_csv(path_spock + '/target_lists/speculoos_target_list_v6.txt', delimiter=' ')
    df_follow_up = pd.read_csv(path_spock + '/target_lists/target_transit_follow_up.txt', delimiter=' ')
    df_special = pd.read_csv(path_spock + '/target_lists/target_list_special.txt', delimiter=' ')

    df_follow_up['nb_hours_surved'] = [0]*len(df_follow_up)
    df_follow_up['nb_hours_threshold'] = [0] * len(df_follow_up)
    df_special['nb_hours_surved'] = [0] * len(df_special)
    df_special['nb_hours_threshold'] = [0] * len(df_special)

    df_pandas = pd.DataFrame({'Sp_ID':pd.concat([df_speculoos['Sp_ID'],df_follow_up['Sp_ID'],df_special['Sp_ID']]),
        'RA': pd.concat([df_speculoos['RA'],df_follow_up['RA'],df_special['RA']]),
        'DEC': pd.concat([df_speculoos['DEC'],df_follow_up['DEC'],df_special['DEC']]),
        'J': pd.concat([df_speculoos['J'],df_follow_up['J'],df_special['J']]),
        'SpT': pd.concat([df_speculoos['SpT'], df_follow_up['SpT'], df_special['SpT']]),
        'nb_hours_surved': pd.concat([df_speculoos['nb_hours_surved'],df_follow_up['nb_hours_surved'],df_special['nb_hours_surved']]),
        'nb_hours_threshold': pd.concat([df_speculoos['nb_hours_threshold'],df_follow_up['nb_hours_threshold'],df_special['nb_hours_threshold']])})
    df_pandas = df_pandas.drop_duplicates()
    df = Table.from_pandas(df_pandas)
    nb_day_date_range = date_range_in_days(date_range)
    doc = Document()
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_beFore = Pt(0)
    par_format.space_after = Pt(6)
    run = par.add_run(observatory.name)
    run.bold = True
    font = run.font
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_beFore = Pt(0)
    par_format.space_after = Pt(12)
    run = par.add_run('Schedule from ' + Time(date_range[0].iso,out_subfmt='date').iso + ' to ' + Time(date_range[1].iso,out_subfmt='date').iso)
    run.bold = True
    font = run.font
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_beFore = Pt(0)
    par_format.space_after = Pt(12)
    run = par.add_run(
        '(Total time = 0hr, technical loss = 0hr, weather loss = 0hr, Exotime = 0hr, cometime = 0hr,   chilean time = 0hr)')
    run.bold = True
    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(255, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_beFore = Pt(0)
    par_format.space_after = Pt(20)
    run = par.add_run(name_operator)
    run.italic = True
    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.space_beFore = Pt(16)
    par_format.space_after = Pt(0)

    for i in range(nb_day_date_range):

        date = date_range[0] + i
        table_schedule = read_night_block(telescope, date)
        sun_set = observatory.sun_set_time(date, which='next').iso
        sun_rise = observatory.sun_rise_time(date, which='next').iso
        moon_illum = int(round(moon_illumination(date) * 100, 0)) * u.percent
        civil_twilights = [Time(observatory.twilight_evening_civil(date, which='next')).iso,
                           Time(observatory.twilight_morning_civil(date + 1, which='nearest')).iso]
        nautic_twilights = [Time(observatory.twilight_evening_nautical(date, which='next')).iso,
                            Time(observatory.twilight_morning_nautical(date + 1, which='nearest')).iso]
        astro_twilights = [Time(observatory.twilight_evening_astronomical(date, which='next')).iso,
                           Time(observatory.twilight_morning_astronomical(date + 1, which='nearest')).iso]
        start_night = table_schedule['start time (UTC)'][0]
        end_night = np.array(table_schedule['end time (UTC)'])[-1]
        night_duration = round((Time(nautic_twilights[1]) - Time(nautic_twilights[0])).jd * 24, 3) * u.hour

        run = par.add_run('Night starting on the ' + Time(date, out_subfmt='date').value)
        run.bold = True
        run.underline = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_beFore = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run('Moon illumination: ' + str(moon_illum))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_beFore = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run(
            'Sunset - Sunrise: ' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.minute) +\
            '  / ' + '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_beFore = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run(
            'Civil/Naut./Astro. twilights: ' + \
            '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.minute)  +\
            '-' + '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.minute)  +\
            ' / ' + '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.minute) +\
            '-' + '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.minute) +\
            '  / ' + '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.minute) +\
            '-' + '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_beFore = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run('Start-end of night (Naut. twil.): ' + '{:02d}'.format(Time(start_night).datetime.hour) + 'h' + '{:02d}'.format(Time(start_night).datetime.minute) +\
                          ' to ' + '{:02d}'.format(Time(end_night).datetime.hour) + 'h' + '{:02d}'.format(Time(end_night).datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_beFore = Pt(0)
        par_format.space_after = Pt(3)
        run = par.add_run('Night duration (Naut. twil.): ' + str(night_duration))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        for j in range(len(table_schedule)):
            # trappist_planets = ['Trappist-1b','Trappist-1c','Trappist-1d','Trappist-1e',
            #                     'Trappist-1f','Trappist-1g','Trappist-1h']
            #
            # if any(table_schedule['target'][j] == p for p in trappist_planets):
            #     idx_target = np.where((df['Sp_ID'] == 'Sp2306-0502'))[0]
            # else:
            if table_schedule['target'][j][-2:] == '_2':
                table_schedule['target'][j] = table_schedule['target'][j][:-2]
            idx_target = np.where((df['Sp_ID'] == table_schedule['target'][j]))[0]

            start_time_target = table_schedule['start time (UTC)'][j]
            end_time_target = table_schedule['end time (UTC)'][j]
            config = table_schedule['configuration'][j]
            try:
                coords = SkyCoord(ra=df['RA'][idx_target].data.data[0] * u.deg, dec=df['DEC'][idx_target].data.data[0] * u.deg)
            except IndexError:
                break
            dist_moon = '34'

            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(0)
            par_format.space_after = Pt(0)
            run = par.add_run(
                'From ' + '{:02d}'.format(Time(start_time_target, out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(start_time_target, out_subfmt='date_hm').datetime.minute) +\
                ' to ' + '{:02d}'.format(Time(end_time_target, out_subfmt='date_hm').datetime.hour) + 'h' + '{:02d}'.format(Time(end_time_target, out_subfmt='date_hm').datetime.minute) +\
                ' : ' + str(table_schedule['target'][j]))
            run.bold = True
            font = run.font
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(0)
            par_format.space_after = Pt(0)
            run = par.add_run('  Note: Prio_target                                         ')
            font = run.font
            font.size = Pt(10)
            font.color.rgb = RGBColor(0, 0, 0)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(0)
            par_format.space_after = Pt(0)
            run = par.add_run(
                '  SPECULOOS : ' + str(df['nb_hours_surved'][idx_target].data.data[0]*u.hour) + ' of obs over ' + str(
                    df['nb_hours_threshold'][idx_target].data.data[0]*u.hour))
            font = run.font
            font.size = Pt(10)
            font.color.rgb = RGBColor(0, 0, 0)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(0)
            par_format.space_after = Pt(0)
            run = par.add_run('Jmag= ' + str(df['J'][idx_target].data.data[0]) + ',  SpT= ' + str(
                df['SpT'][idx_target].data[0]))  # + ', Moon at ' + str(dist_moon))
            font = run.font
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(0)
            par_format.space_after = Pt(3)
            run = par.add_run(' RA = ' + str('{:02d}'.format(int(coords.ra.hms[0]))) + " " + str('{:02d}'.format(int(coords.ra.hms[1]))) + " " + str('{:05.3f}'.format(round(coords.ra.hms[2], 3))) + \
                ', DEC = ' + str('{:02d}'.format(int(coords.dec.dms[0]))) + " " + str('{:02d}'.format(int(abs(coords.dec.dms[1])))) + " " + str('{:05.3f}'.format(round(abs(coords.dec.dms[2]), 3))) + ', ' + str(config[2:-2]).replace('\'',' '))
            font = run.font
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_beFore = Pt(16)
            par_format.space_after = Pt(0)

    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    if telescope == 'TS_La_Silla':
        doc.save(path_spock + '/TRAPPIST_schedules_docx/TS_' + Time(date_range[0],out_subfmt='date').value.replace('-','') + '_to_' +\
                 Time(date_range[1],out_subfmt='date').value .replace('-','')  +'.docx')
    if telescope == 'TN_Oukaimeden':
        doc.save(path_spock + '/TRAPPIST_schedules_docx/TN_' + Time(date_range[0],out_subfmt='date').value.replace('-','') + '_to_' +\
                 Time(date_range[1],out_subfmt='date').value .replace('-','')  +'.docx')

def date_range_in_days(date_range):
    date_format = "%Y-%m-%d %H:%M:%S.%f"
    date_start = datetime.strptime(date_range[0].value, date_format)
    date_end = datetime.strptime(date_range[1].value, date_format)
    date_range_in_days = (date_end - date_start).days
    return date_range_in_days
