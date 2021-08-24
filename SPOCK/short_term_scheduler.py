#!/anaconda3/bin/python3.6
from astropy.time import Time
from astropy.table import unique, Table, vstack
from astropy import units as u
from astropy.coordinates import SkyCoord, EarthLocation, AltAz
from astroplan.utils import time_grid_from_range
from astroplan import FixedTarget, AltitudeConstraint, MoonSeparationConstraint,AtNightConstraint,AirmassConstraint,\
    TimeConstraint, observability_table
from astroplan.plots import dark_style_sheet,plot_airmass
from astroplan import Observer,moon_illumination, is_observable
from astroplan.periodic import EclipsingSystem
from astroplan.constraints import is_event_observable
from colorama import Fore
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import *
from datetime import date, timedelta, datetime
from eScheduler.spe_schedule import SPECULOOSScheduler, Schedule, ObservingBlock,Transitioner
import gspread
import matplotlib.pyplot as plt
import numpy as np
from oauth2client.service_account import ServiceAccountCredentials
import os
import pandas as pd
import requests
from .upload_night_plans import upload_np_artemis, upload_np_saint_ex, upload_np_io, upload_np_gany, upload_np_euro,\
    upload_np_calli, upload_np_tn, upload_np_ts
from .make_night_plans import make_np
import subprocess
import sys
import shutil
import SPOCK.ETC as ETC
from SPOCK import user_portal, pwd_portal, pwd_appcs, path_spock, path_credential_json, target_list_from_stargate_path


scope = ['https://spreadsheets.google.com/feeds',
         'https://www.googleapis.com/auth/drive']
creds = ServiceAccountCredentials.from_json_keyfile_name(path_credential_json, scope)
client = gspread.authorize(creds)

dt = Time('2018-01-02 00:00:00',scale='tcg')-Time('2018-01-01 00:00:00',scale='tcg')  # 1 day
constraints = [AltitudeConstraint(min=24*u.deg), AtNightConstraint()]  # MoonSeparationConstraint(min=30*u.deg)


def charge_observatories(name):
    """

    Parameters
    ----------
    name: name of the observatory  "SSO" or "SNO" or "Saint-ex" or "TN_Oukaimeden or "TS_La_Silla"

    Returns
    -------
    astroplan.Observer object with name and coordinates of the observatory

    """
    observatories = []
    # Observatories
    if 'SSO' in str(name):
        location = \
            EarthLocation.from_geodetic(-70.40300000000002*u.deg, -24.625199999999996*u.deg,2635.0000000009704*u.m)
        observatories.append(Observer(location=location, name="SSO", timezone="UTC"))

    if 'SNO' in str(name):
        location_SNO = EarthLocation.from_geodetic(-16.50583131*u.deg, 28.2999988*u.deg, 2390*u.m)
        observatories.append(Observer(location=location_SNO, name="SNO", timezone="UTC"))

    if 'Saint-Ex' in str(name):
        location_saintex = \
            EarthLocation.from_geodetic(-115.48694444444445*u.deg, 31.029166666666665*u.deg, 2829.9999999997976*u.m)
        observatories.append(Observer(location=location_saintex, name="Saint-Ex", timezone="UTC"))

    if 'TS_La_Silla' in str(name):
        location_TSlasilla = \
            EarthLocation.from_geodetic(-70.73000000000002*u.deg, -29.25666666666666*u.deg, 2346.9999999988418*u.m)
        observatories.append(Observer(location=location_TSlasilla, name="TS_La_Silla", timezone="UTC"))

    if 'TN_Oukaimeden' in str(name):
        location_TNOuka = EarthLocation.from_geodetic( -7.862263*u.deg, 31.20516*u.deg,2751*u.m)
        observatories.append(Observer(location=location_TNOuka, name="TN_Oukaimeden", timezone="UTC"))

    if 'Munich' in str(name):
        location_munich = EarthLocation.from_geodetic(48.2*u.deg, -11.6*u.deg, 600*u.m)
        observatories.append(Observer(location=location_munich, name="Munich", timezone="UTC"))

    return observatories


def target_list_good_coord_format(df=None, path_target_list=None):

    """
    Give target corrdinates in ICRS format (used for astropy.coordinates SkyCoord function)

    Parameters
    ----------
    path_target_list: path on your computer toward the target list, by default take the one on the Cambridge server
    df: panda dataframe of the target list

    Returns
    -------
    targets: targets list with the following format : [<FixedTarget "Sp0002+0115" at SkyCoord (ICRS): (ra, dec) in deg (0.52591667, 1.26003889)>,


    """
    if df is None:
        df = pd.read_csv(path_target_list, delimiter=' ')
    target_table_spc = Table.from_pandas(df)
    targets = [FixedTarget(coord=SkyCoord(ra=target_table_spc['RA'][i]* u.degree,
                                          dec=target_table_spc['DEC'][i] * u.degree),
                           name=target_table_spc['Sp_ID'][i]) for i in range(len(target_table_spc['RA']))]
    return targets


class Schedules:
    """
    Class Schedules to make night plans
    """

    def __init__(self):
        self.altitude_constraint = 25
        self.target_list = None
        self.telescopes = []
        self.telescope = []
        self.start_end_range = None  # date_range
        self.day_of_night = None
        self.constraints = [AtNightConstraint.twilight_civil()]
        self.moon_constraint = 30
        self.observatory = None
        self.observatory_name = None
        self.SS1_night_blocks = None
        self.scheduled_table = None
        self.SS1_night_blocks_old = None
        self.scheduled_table_sorted = None
        self.targets_follow_up = None
        self.target_table_spc_follow_up = None
        self.target_list_follow_up = None
        self.target_list_special = None
        self.targets = None
        self.target_table_spc = []

    @property
    def start_of_observation(self):
        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')  # 1 day
        start_between_civil_nautical = Time((Time(
            self.observatory.twilight_evening_nautical(self.day_of_night, which='next')).value +
                                             Time(self.observatory.twilight_evening_civil(
                                                 self.day_of_night, which='next')).value) / 2, format='jd')
        return start_between_civil_nautical

    @property
    def end_of_observation(self):
        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')  # 1 day
        end_between_nautical_civil = Time((Time(
            self.observatory.twilight_morning_nautical(self.day_of_night + 1, which='nearest')).value +
                                           Time(self.observatory.twilight_morning_civil(
                                               self.day_of_night + 1, which='nearest')).value) / 2, format='jd')
        return end_between_nautical_civil

    def load_parameters(self, filename_list_special=None, filename_follow_up=None):
        """

        Parameters
        ----------
        filename_list_special: name of the file with special targets (stars not in SPECULOOS survey) to observe
        filename_follow_up: name of the file with follow up candidates targets (planets with ephemeris)

        Returns
        -------

        """

        if (filename_list_special is None) or (filename_follow_up is None):
            sh = client.open('SPECULOOS WG6')
            # Read stars lists
            worksheet_follow_up = sh.worksheet("Annex_Targets_V1-PLANETS")
            dataframe = pd.DataFrame(worksheet_follow_up.get_all_records())
            self.target_table_spc_follow_up = dataframe.rename(columns={"sp_id": "Sp_ID", "gaia_dr2": "Gaia_ID",
                                                                        "period": "P", "period_e": "P_err",
                                                                        "duration": "W", "duration_e": "W_err",
                                                                        "dec": "DEC", "ra": "RA",
                                                                        "dec_err": "DEC_err", "ra_err": "RA_err",

                                                                        })
            self.targets_follow_up = target_list_good_coord_format(df=self.target_table_spc_follow_up)

            # Read follow up (planet candidates) list
            self.target_table_spc_follow_up['W'] /= 24
            self.target_table_spc_follow_up['W_err'] /= 24
            worksheet_special = sh.worksheet("Annex_Targets_V2-STARS")
            dataframe = pd.DataFrame(worksheet_special.get_all_records())
            self.target_table_spc = dataframe.rename(columns={"spc": "Sp_ID", "gaia": "Gaia_ID", "dec": "DEC",
                                                              "ra": "RA", "dec_err": "DEC_err", "ra_err": "RA_err",
                                                              "mag_j": "J", "V_mag": "V"})
            self.targets = target_list_good_coord_format(df=self.target_table_spc)

        if filename_list_special is not None:
            self.target_list_special = filename_list_special
            self.targets = target_list_good_coord_format(path_target_list=filename_list_special)
            self.target_table_spc = Table.from_pandas(pd.read_csv(filename_list_special, delimiter=' '))

        if filename_follow_up is not None:
            self.target_list_follow_up = filename_follow_up
            self.targets_follow_up = target_list_good_coord_format(path_target_list=filename_follow_up)
            self.target_table_spc_follow_up = Table.from_pandas(pd.read_csv(filename_follow_up, delimiter=' '))

    def night_duration(self, day):
        """

        Parameters
        ----------
        day: day str format '%y%m%d HH:MM:SS.sss'

        Returns
        -------
        duration of the day in astropy Time format
        """
        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')
        return self.end_of_observation - self.start_of_observation

    def monitoring(self, input_name, airmass_max, time_monitoring):
        """
        Function to add a monitoring night block in  the plans
        Parameters
        ----------
        input_name: name of the target to monitor
        airmass_max: maximum airmass for the monitoring
        time_monitoring: time range of the monitoring in minutes

        Returns
        -------
        Create night block

        """
        idx_first_target = int(np.where((self.target_table_spc['Sp_ID'] == input_name))[0])
        dur_mon_target = time_monitoring * u.minute
        constraints_monitoring_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                         MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                         AirmassConstraint(max=airmass_max, boolean_constraint=True),
                                         TimeConstraint((
                                             Time(self.observatory.twilight_evening_nautical(self.day_of_night,
                                                                                             which='next'))),
                                             (Time(self.observatory.twilight_morning_nautical(self.day_of_night+1,
                                                                                              which='nearest'))))]
        if self.target_table_spc['texp_spc'][idx_first_target] == 0:
            self.target_table_spc['texp_spc'][idx_first_target], self.target_table_spc['Filter_spc'][idx_first_target] = \
                self.exposure_time(input_name=self.target_table_spc['Sp_ID'][idx_first_target],
                                   target_list=self.target_table_spc)
        blocks = []
        a = ObservingBlock(self.targets[idx_first_target], dur_mon_target, -1,
                           constraints=constraints_monitoring_target,
                           configuration={"filt": str(self.target_table_spc['Filter_spc'][idx_first_target]),
                                          "texp":  str(self.target_table_spc['texp_spc'][idx_first_target])})
        blocks.append(a)
        transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
        seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night+1)
        sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_monitoring_target,
                                                  observer=self.observatory, transitioner=transitioner)
        sequen_scheduler_ss1(blocks, seq_schedule_ss1)
        self.SS1_night_blocks = seq_schedule_ss1.to_table()
        return self.SS1_night_blocks

    def dome_rotation(self):
        """
        Function to add a dome rotation night block in the plans
        Returns
        -------
        Create a night block in between two targets or in the middle of  the night to ensure th dome is doing a full
        rotation

        """
        # not necessary
        dur_dome_rotation = 5 / 60 / 24 * u.day  # 5min

        sun_set = Time(self.observatory.twilight_evening_nautical(self.day_of_night, which='next'))
        sun_rise = Time(self.observatory.twilight_morning_nautical(self.day_of_night+1, which='nearest'))

        time_resolution = 1*u.hour
        time_grid = time_grid_from_range([sun_set, sun_rise], time_resolution=time_resolution)
        dom_rot_possible = False

        self.make_scheduled_table()
        end_times = Time(self.scheduled_table['end time (UTC)'][:]).iso

        for i in range(1, len(time_grid)-1):

            start = time_grid[i]
            # print(start.iso)
            end = start + dur_dome_rotation
            idx = np.where((start < end_times[:]))[0]

            coords = SkyCoord(str(int(self.scheduled_table['ra (h)'][idx][0])) + 'h' +
                              str(int(self.scheduled_table['ra (m)'][idx][0])) + 'm' +
                              str(round(self.scheduled_table['ra (s)'][idx][0], 3)) + 's' + ' ' +
                              str(int(self.scheduled_table['dec (d)'][idx][0])) + 'd' +
                              str(abs(int(self.scheduled_table['dec (m)'][idx][0]))) + 'm' +
                              str(abs(round(self.scheduled_table['dec (s)'][idx][0], 3))) + 's').transform_to(
                AltAz(obstime=start,location=self.observatory.location))
            coords_dome_rotation = SkyCoord(alt=coords.alt, az=(coords.az.value - 180) * u.deg, obstime=start,
                                            frame='altaz', location=self.observatory.location)
            if (coords.alt.value > 60) or (coords.alt.value < 30):
                dom_rot_possible = False
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
                      ' dome rotation not possible at that time because of altitude constraint')

            else:
                target = FixedTarget(coord=SkyCoord(ra=coords_dome_rotation.icrs.ra.value * u.degree,
                                                    dec=coords_dome_rotation.icrs.dec.value * u.degree),
                        name='dome_rot')
                dom_rot_possible = True

                df = pd.DataFrame({'target': target.name, 'start time (UTC)': start.iso, 'end time (UTC)': end.iso,
                                   'duration (minutes)': dur_dome_rotation.value * 24 * 60,
                                   'ra (h)': target.coord.ra.hms[0],
                                   'ra (m)': target.coord.ra.hms[1], 'ra (s)': target.coord.ra.hms[2],
                                   'dec (d)': target.coord.dec.dms[0], 'dec (m)': target.coord.dec.dms[1],
                                   'dec (s)': target.coord.dec.dms[2], 'configuration': "{\'filt=I+z\', \'texp=10\'}"},
                                  index=[0])
                self.SS1_night_blocks = Table.from_pandas(df)
                return self.SS1_night_blocks

                break

        if not dom_rot_possible:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No Dom rotation possible that night')

    def special_target_with_start_end(self, input_name):
        """
        Function to add a special target night block in the plans with specific start/end times
        from self.start_end_range
        Parameters
        ----------
        input_name: name of the special target to schedule

        Returns
        -------
        Create a night block with special target observed in the given time range (self.start_end_range)

        """
        self.observatory = charge_observatories(self.observatory_name)[0]
        start = self.start_end_range[0]
        end = self.start_end_range[1]

        if (start >= self.end_of_observation) or \
                (end <= self.start_of_observation):
            sys.exit(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + 'Start time (or End time) is not on the same day.')

        dur_obs_both_target = (self.night_duration(self.day_of_night) / (2 * u.day)) * 2 * u.day
        constraints_special_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                      MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                      TimeConstraint(start, end)]
        idx_to_insert_target = int(np.where((self.target_table_spc['Sp_ID'] == input_name))[0])
        if self.target_table_spc['texp_spc'][idx_to_insert_target] == 0 \
                or self.target_table_spc['texp_spc'][idx_to_insert_target] == "00":
            self.target_table_spc['texp_spc'][idx_to_insert_target] = \
                self.target_table_spc['texp_spc'][idx_to_insert_target], \
                self.target_table_spc['Filter_spc'][idx_to_insert_target] =\
                self.exposure_time(input_name=self.target_table_spc['Sp_ID'][idx_to_insert_target],
                                   target_list=self.target_table_spc)
        observable = is_observable(constraints_special_target, self.observatory, self.targets[idx_to_insert_target],
                                   time_range=(start, end))
        if observable:
            blocks = []
            a = ObservingBlock(self.targets[idx_to_insert_target], dur_obs_both_target, -1,
                               constraints=constraints_special_target,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][idx_to_insert_target]),
                                              "texp":  str(self.target_table_spc['texp_spc'][idx_to_insert_target])})
            blocks.append(a)
            transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
            seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
            sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_special_target, observer=self.observatory,
                                                      transitioner=transitioner)
            sequen_scheduler_ss1(blocks, seq_schedule_ss1)
            self.SS1_night_blocks = seq_schedule_ss1.to_table()
            return self.SS1_night_blocks
        else:
            sys.exit(Fore.RED + 'ERROR: ' + Fore.BLACK
                     + " Observation impossible due to unrespected altitude and/or moon constraints. ")


    def special_target(self, input_name):
        """
        Function to add a special target night block in the plans
        Parameters
        ----------
        input_name: name of the special target to schedule

        Returns
        -------
        Create a night block with special target observed as much as possible given the constraints

        """
        self.observatory = charge_observatories(self.observatory_name)[0]
        dur_obs_both_target = (self.night_duration(self.day_of_night)/(2*u.day))*2*u.day
        constraints_special_target = [AltitudeConstraint(min=self.altitude_constraint*u.deg),
                                      MoonSeparationConstraint(min=self.moon_constraint*u.deg),
                                      TimeConstraint(self.start_of_observation,self.end_of_observation)]
        idx_first_target = int(np.where((self.target_table_spc["Sp_ID"] == input_name))[0])
        if int(self.target_table_spc['texp_spc'][idx_first_target]) == 0:
            self.target_table_spc['texp_spc'][idx_first_target], self.target_table_spc['Filter_spc'][idx_first_target] = \
                self.exposure_time(input_name=self.target_table_spc['Sp_ID'][idx_first_target],
                                   target_list=self.target_table_spc)
        blocks = []
        a = ObservingBlock(self.targets[idx_first_target], dur_obs_both_target, -1,
                           constraints=constraints_special_target,
                           configuration={"filt": str(self.target_table_spc['Filter_spc'][idx_first_target]),
                                          "texp":  str(self.target_table_spc['texp_spc'][idx_first_target])})
        blocks.append(a)
        transitioner = Transitioner(slew_rate=11*u.deg/u.second)
        seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night+1)
        sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_special_target,
                                                  observer=self.observatory, transitioner=transitioner)
        sequen_scheduler_ss1(blocks,seq_schedule_ss1)
        self.SS1_night_blocks=seq_schedule_ss1.to_table()
        if len(self.SS1_night_blocks) == 0:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + 'Impossible to schedule target ' + input_name +
                  ' at this time range and/or with those constraints')
        return self.SS1_night_blocks

    def transit_follow_up(self, input_name):
        """
        Function to add a  night block with a transit of a follow candidate in the plans
        Parameters
        ----------
        input_name: name of the follow up candidate to schedule

        Returns
        -------
        Create a night block with transiting candidate observed for: transit duration * (1 + 1.5) centre on T0_predicted

        """
        self.observatory = charge_observatories(self.observatory_name)[0]

        constraints_follow_up = [AltitudeConstraint(min=self.altitude_constraint * u.deg), AtNightConstraint(),
                                 TimeConstraint(Time(self.day_of_night), Time(self.day_of_night+1))]
        df = self.target_table_spc_follow_up

        if input_name is not None:
            i = int(np.where((df['Sp_ID'] == input_name))[0])
            blocks = []
            epoch = Time(df['T0'][i], format='jd')
            period = df['P'][i] * u.day
            duration = df['W'][i] * u.day
            oot_time = duration.value * 1.5 * u.day
            T0_err_transit = df['T0_err'][i]
            P_err_transit = df['P_err'][i]
            W_err_transit = df['W_err'][i]
            target_transit = EclipsingSystem(primary_eclipse_time=epoch, orbital_period=period, duration=duration,
                                             name=df['Sp_ID'][i])
            print(Fore.GREEN + Fore.GREEN + Fore.GREEN + 'INFO: ' + Fore.BLACK + ' ' + Fore.BLACK +
                  Fore.BLACK + str(df['Sp_ID'][i]) + ' next transit: ',
                  Time(target_transit.next_primary_eclipse_time(self.day_of_night, n_eclipses=1)).iso)
            timing_to_obs_jd = Time(target_transit.next_primary_eclipse_time(self.day_of_night, n_eclipses=1)).jd
            n_transits = 1
            try:
                ing_egr = target_transit.next_primary_ingress_egress_time(self.day_of_night, n_eclipses=n_transits)
            except ValueError:
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No transit of ', df['Sp_ID'][i],
                      ' on the period chosen')

            observable = is_event_observable(constraints_follow_up, self.observatory, self.targets_follow_up,
                                             times_ingress_egress=ing_egr)
            if np.any(observable):
                err_T0_neg = T0_err_transit  # timing_to_obs_jd[0] -
                # (np.round((timing_to_obs_jd[0] - epoch.jd) / period.value, 1) *(period.value -
                # P_err_transit) + (epoch.jd - T0_err_transit))
                err_T0_pos = T0_err_transit  #(np.round((timing_to_obs_jd[0] - epoch.jd) / period.value, 1) *
                # (period.value + P_err_transit) + (epoch.jd + T0_err_transit)) - timing_to_obs_jd[0]
                start_transit = Time(ing_egr[0][0].value - err_T0_neg - oot_time.value - W_err_transit,
                                     format='jd')
                end_transit = Time(ing_egr[0][1].value + err_T0_pos + oot_time.value + W_err_transit,
                                   format='jd')
                dur_obs_transit_target = (end_transit - start_transit).value * 1. * u.day
                if (end_transit > self.end_of_observation) \
                        or (start_transit < self.start_of_observation):
                    if (Time(ing_egr[0][1]) < self.end_of_observation) \
                            and (Time(ing_egr[0][0]) > self.start_of_observation):
                        constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                      TimeConstraint(start_transit, end_transit),
                                                      MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                                      AtNightConstraint(max_solar_altitude=-12 * u.deg)]
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ', str(df['Sp_ID'][i]),
                              ' : ', start_transit.iso)
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' end_transit of ', str(df['Sp_ID'][i]),
                              ' : ', end_transit.iso)
                        print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Time out of transit not optimal.')
                        idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                        if df['texp_spc'][idx_first_target] == 0:
                            df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = self.exposure_time(
                                input_name=df['Sp_ID'][idx_first_target], target_list=self.target_table_spc_follow_up,
                                day=self.day_of_night)
                        a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                           constraints=constraints_transit_target,
                                           configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                          "texp":  str(df['texp_spc'][idx_first_target])})
                        blocks.append(a)
                        transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                        seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
                        sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                                  observer=self.observatory,
                                                                  transitioner=transitioner)
                        sequen_scheduler_ss1(blocks, seq_schedule_ss1)

                    else:
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ', str(df['Sp_ID'][i]),
                              ' : ', Time(ing_egr[0][0]).iso)
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' end_transit of ', str(df['Sp_ID'][i]),
                              ' : ', Time(ing_egr[0][1]).iso)
                        print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Transit not full.')
                        constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                      MoonSeparationConstraint(min= self.moon_constraint* u.deg),
                                                      TimeConstraint(Time(ing_egr[0][0]), Time(ing_egr[0][1]))]
                        idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                        if df['texp_spc'][idx_first_target] == 0:
                            df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = self.exposure_time(
                                input_name=df['Sp_ID'][idx_first_target], target_list=self.target_table_spc_follow_up)
                        a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                           constraints=constraints_transit_target,
                                           configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                          "texp":  str(df['texp_spc'][idx_first_target])})
                        blocks.append(a)
                        transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                        seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
                        sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                                  observer=self.observatory,
                                                                  transitioner=transitioner)
                        sequen_scheduler_ss1(blocks, seq_schedule_ss1)

                else:
                    constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                  MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                                  TimeConstraint(start_transit, end_transit)]
                    idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                    if df['texp_spc'][idx_first_target] == 0:
                        df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = \
                            self.exposure_time(input_name=df['Sp_ID'][idx_first_target],
                                               target_list=self.target_table_spc_follow_up)
                    a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                       constraints=constraints_transit_target,
                                       configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                      "texp": str(df['texp_spc'][idx_first_target])})
                    blocks.append(a)
                    transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                    seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
                    sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                              observer=self.observatory,
                                                              transitioner=transitioner)
                    sequen_scheduler_ss1(blocks, seq_schedule_ss1)
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ', str(df['Sp_ID'][i]),
                          ' : ', Time(ing_egr[0][0]).iso)
                    print(Fore.GREEN + Fore.GREEN + 'INFO: ' + Fore.BLACK + ' ' + Fore.BLACK + ' end_transit of ',
                          str(df['Sp_ID'][i]), ' : ', Time(ing_egr[0][1]).iso)
                    print(Fore.GREEN + Fore.GREEN + 'INFO: ' + Fore.BLACK + ' ' + Fore.BLACK +
                          ' Transit is expected to be full.')
                    if len(seq_schedule_ss1.to_table()['target']) == 0:
                        print(Fore.RED + 'ERROR:  ' + Fore.BLACK +
                              ' The moon is too closed for the transit to be observed')

                if len(seq_schedule_ss1.to_table()['target']) != 0:
                    self.SS1_night_blocks = seq_schedule_ss1.to_table()
                    return self.SS1_night_blocks
            else:
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transit of ', df['Sp_ID'][i], ' this day')

        else:
            for i in range(len(df['Sp_ID'])):
                blocks = []
                epoch = Time(df['T0'][i], format='jd')
                period = df['P'][i] * u.day
                duration = df['W'][i] * u.day
                oot_time = duration.value * 1.5 * u.day
                T0_err_transit = df['T0_err'][i]
                P_err_transit = df['P_err'][i]
                W_err_transit = df['W_err'][i]
                target_transit = EclipsingSystem(primary_eclipse_time=epoch, orbital_period=period,
                                                 duration=duration, name=df['Sp_ID'][i])
                print(Fore.GREEN + Fore.GREEN + Fore.GREEN + 'INFO: ' + Fore.BLACK + ' ' +
                      Fore.BLACK + Fore.BLACK + str(df['Sp_ID'][i]) + ' next transit: ',
                      Time(target_transit.next_primary_eclipse_time(self.day_of_night, n_eclipses=1)).iso)
                timing_to_obs_jd = Time(target_transit.next_primary_eclipse_time(self.day_of_night, n_eclipses=1)).jd
                n_transits = 1
                try:
                    ing_egr = target_transit.next_primary_ingress_egress_time(self.day_of_night, n_eclipses=n_transits)
                except ValueError:
                    print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No transit of ',
                          df['Sp_ID'][i], ' on the period chosen')

                observable = is_event_observable(constraints,self.observatory,
                                                 self.targets_follow_up, times_ingress_egress=ing_egr)
                if np.any(observable):
                    err_T0_neg = timing_to_obs_jd[0] - (np.round((timing_to_obs_jd[0] - epoch.jd) / period.value, 1) *
                                                        (period.value - P_err_transit) + (epoch.jd - T0_err_transit))
                    err_T0_pos = (np.round((timing_to_obs_jd[0] - epoch.jd) / period.value, 1) *
                                  (period.value + P_err_transit) + (epoch.jd + T0_err_transit)) - timing_to_obs_jd[0]
                    start_transit = Time(ing_egr[0][0].value - err_T0_neg - oot_time.value - W_err_transit, format='jd')
                    end_transit = Time(ing_egr[0][1].value + err_T0_pos + oot_time.value + W_err_transit, format='jd')
                    dur_obs_transit_target = (end_transit - start_transit).value * 1. * u.day
                    if (end_transit > self.end_of_observation) \
                            or (start_transit < self.start_of_observation):

                        if (Time(ing_egr[0][1]) < self.end_of_observation)\
                                and (Time(ing_egr[0][0]) > self.start_of_observation):

                            constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                          TimeConstraint(start_transit, end_transit),
                                                          MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                                          AtNightConstraint(max_solar_altitude=-12*u.deg)]
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ',
                                  str(df['Sp_ID'][i]), ' : ', start_transit.iso)
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' end_transit of ',
                                  str(df['Sp_ID'][i]), ' : ', end_transit.iso)
                            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Time out of transit not optimal.')
                            idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                            if df['texp_spc'][idx_first_target] == 0:
                                df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = self.exposure_time(
                                    input_name=df['Sp_ID'][idx_first_target],
                                    target_list=self.target_table_spc_follow_up)
                            a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                               constraints=constraints_transit_target,
                                               configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                              "texp":  str(df['texp_spc'][idx_first_target])})
                            blocks.append(a)
                            transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                            seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
                            sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                                      observer=self.observatory,
                                                                      transitioner=transitioner)
                            sequen_scheduler_ss1(blocks, seq_schedule_ss1)

                        else:
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ',
                                  str(df['Sp_ID'][i]), ' : ', Time(ing_egr[0][0]).iso)
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' end_transit of ',
                                  str(df['Sp_ID'][i]), ' : ', Time(ing_egr[0][1]).iso)
                            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Transit not full.')
                            constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                          MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                                          TimeConstraint(Time(ing_egr[0][0]), Time(ing_egr[0][1]))]
                            # continue
                            idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                            if df['texp_spc'][idx_first_target] == 0:
                                df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = self.exposure_time(
                                    input_name=df['Sp_ID'][idx_first_target],
                                    target_list=self.target_table_spc_follow_up)
                            a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                               constraints=constraints_transit_target,
                                               configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                              "texp":  str(df['texp_spc'][idx_first_target])})
                            blocks.append(a)
                            transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                            seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night + 1)
                            sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                                      observer=self.observatory,
                                                                      transitioner=transitioner)
                            sequen_scheduler_ss1(blocks, seq_schedule_ss1)

                    else:
                        constraints_transit_target = [AltitudeConstraint(min=self.altitude_constraint * u.deg),
                                                      MoonSeparationConstraint(min=self.moon_constraint * u.deg),
                                                      TimeConstraint(start_transit, end_transit)]
                        idx_first_target = int(np.where((df['Sp_ID'] == df['Sp_ID'][i]))[0])
                        if df['texp_spc'][idx_first_target] == 0:
                            df['texp_spc'][idx_first_target], df['Filter_spc'][idx_first_target] = \
                                self.exposure_time(input_name=df['Sp_ID'][idx_first_target],
                                                   target_list=self.target_table_spc_follow_up)
                        a = ObservingBlock(self.targets_follow_up[idx_first_target], dur_obs_transit_target, -1,
                                           constraints=constraints_transit_target,
                                           configuration={"filt": str(df['Filter_spc'][idx_first_target]),
                                                          "texp":  str(df['texp_spc'][idx_first_target])})
                        blocks.append(a)
                        transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
                        seq_schedule_ss1 = Schedule(self.day_of_night, self.day_of_night+1)
                        sequen_scheduler_ss1 = SPECULOOSScheduler(constraints=constraints_transit_target,
                                                                  observer=self.observatory,
                                                                  transitioner=transitioner)
                        sequen_scheduler_ss1(blocks, seq_schedule_ss1)
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' start_transit of ', str(df['Sp_ID'][i]),
                              ' : ', Time(ing_egr[0][0]).iso)
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' end_transit of ', str(df['Sp_ID'][i]),
                              ' : ', Time(ing_egr[0][1]).iso)
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Transit is expected to be full.')
                        if len(seq_schedule_ss1.to_table()['target']) == 0:
                            print(Fore.RED + 'ERROR:  ' + Fore.BLACK +
                                  ' The moon is too closed for the transit to be observed')

                    if len(seq_schedule_ss1.to_table()['target']) != 0:
                        self.SS1_night_blocks = seq_schedule_ss1.to_table()
                        return self.SS1_night_blocks
                else:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transit of ', df['Sp_ID'][i], ' this day')

    def locking_observations(self):
        night_block = self.SS1_night_blocks.to_pandas()

        night_block.to_csv(path_spock + '/DATABASE/' + self.telescope + '/' +
                           'Locked_obs/'+'lock_night_block_'+self.telescope+'_' + 
                           Time(self.day_of_night, out_subfmt='date').iso+'.txt', index=None)

        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + 'Observation block saved as ' +
              path_spock + '/DATABASE/' + self.telescope + '/' +
                           'Locked_obs/'+'lock_night_block_'+self.telescope+'_' + 
              Time(self.day_of_night, out_subfmt='date').iso+'.txt')

    def planification(self):
        end_scheduled_table = pd.DataFrame(columns=['target', 'start time (UTC)', 'end time (UTC)',
                                                    'duration (minutes)', 'ra (h)', 'ra (m)', 'ra (s)',
                                                    'dec (d)', 'dec (m)', 'dec (s)', 'configuration'])
        try:
            self.SS1_night_blocks['target'][0]
        except TypeError:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No block to insert ')
        if self.SS1_night_blocks['target'][0] in self.scheduled_table['target']:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' This target is already scheduled this day')
        for i in range(len(self.scheduled_table['target'])):
            # print(i)
            try:
                self.SS1_night_blocks['target'][0]
            except IndexError:
                sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No block to insert ')
            except TypeError:
                sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No block to insert ')
            end_before_cut = self.scheduled_table['end time (UTC)'][i]
            start_before_cut = self.scheduled_table['start time (UTC)'][i]

            if self.SS1_night_blocks is None:
                sys.exit('WARNING : No block to insert !')

            if not (self.scheduled_table_sorted is None):
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Several transits this night')
                if (np.any(str(self.SS1_night_blocks['target'][0]).find('Trappist-1')) == 0) and \
                        (np.any(str(self.scheduled_table_sorted['target'][0]).find('Trappist-1')) == 0):
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' two TRAPPIST-1 planets this night!')
                    self.SS1_night_blocks['start time (UTC)'][0] = \
                        min(self.SS1_night_blocks['start time (UTC)'][0],
                            self.scheduled_table_sorted['start time (UTC)'][0])
                    self.SS1_night_blocks['end time (UTC)'][0] = \
                        max(self.SS1_night_blocks['end time (UTC)'][0],
                            self.scheduled_table_sorted['end time (UTC)'][0])
                    self.SS1_night_blocks['duration (minutes)'][0] = \
                        (Time(self.SS1_night_blocks['end time (UTC)'][0]) -
                         Time(self.SS1_night_blocks['start time (UTC)'][0])).value*24*60
                    end_scheduled_table.append({'target': self.SS1_night_blocks['target'][0],
                                                'start time (UTC)': self.SS1_night_blocks['start time (UTC)'][0],
                                                'end time (UTC)': self.SS1_night_blocks['end time (UTC)'][0],
                                                'duration (minutes)': self.SS1_night_blocks['duration (minutes)'][0],
                                                'ra (h)': self.SS1_night_blocks['ra (h)'][0],
                                                'ra (m)': self.SS1_night_blocks['ra (m)'][0],
                                                'ra (s)': self.SS1_night_blocks['ra (s)'][0],
                                                'dec (d)': self.SS1_night_blocks['dec (d)'][0],
                                                'dec (m)': self.SS1_night_blocks['dec (m)'][0],
                                                'dec (s)': self.SS1_night_blocks['dec (s)'][0],
                                                'configuration': self.SS1_night_blocks['configuration'][0]},
                                               ignore_index=True)
            # situation 1
            if (self.SS1_night_blocks['start time (UTC)'][0] <= 
                    self.start_of_observation.iso):
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 1')
                self.SS1_night_blocks['start time (UTC)'][0] = \
                    self.start_of_observation.iso

            if (self.SS1_night_blocks['start time (UTC)'][0] < start_before_cut) and \
                    (self.SS1_night_blocks['start time (UTC)'][0] < end_before_cut):

                # situation 2
                if (self.SS1_night_blocks['end time (UTC)'][0] > start_before_cut) and \
                        (self.SS1_night_blocks['end time (UTC)'][0] < end_before_cut): 
                    # case 2 # if new_block ends beFore the end of existing block
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 2')
                    self.scheduled_table['start time (UTC)'][i] = self.SS1_night_blocks['end time (UTC)'][0]
                    self.scheduled_table['duration (minutes)'][i] = \
                        (Time(self.scheduled_table['end time (UTC)'][i]) -
                         Time(self.scheduled_table['start time (UTC)'][i])).value*24*60

                # situation 3
                if self.SS1_night_blocks['end time (UTC)'][0] <= start_before_cut:
                    self.scheduled_table['duration (minutes)'][i] = \
                        (Time(self.scheduled_table['end time (UTC)'][i]) -
                         Time(self.scheduled_table['start time (UTC)'][i])).value*24*60
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 3, no change made to initial schedule')

                # situation 4
                elif (self.SS1_night_blocks['end time (UTC)'][0] >= end_before_cut) and \
                        (self.SS1_night_blocks['end time (UTC)'][0] <= 
                         self.end_of_observation.iso):
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 4')
                    self.scheduled_table[i] = self.SS1_night_blocks[0]  # a way to erase self.scheduled_table block

                # situation 5
                elif (self.SS1_night_blocks['end time (UTC)'][0] >= 
                      self.end_of_observation.iso):
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 5')
                    self.SS1_night_blocks['end time (UTC)'][0] = \
                        self.end_of_observation.iso
                    self.scheduled_table[i] = self.SS1_night_blocks[0]  # a way to erase self.scheduled_table block

            if self.SS1_night_blocks['start time (UTC)'][0] >= start_before_cut:

                # situation 6
                if (self.SS1_night_blocks['start time (UTC)'][0] <= end_before_cut) and \
                        (self.SS1_night_blocks['end time (UTC)'][0] <= end_before_cut):
                    
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 6')
                    self.scheduled_table['end time (UTC)'][i] = self.SS1_night_blocks['start time (UTC)'][0]
                    self.scheduled_table['duration (minutes)'] = \
                        (Time(self.scheduled_table['end time (UTC)'][i]) - 
                         Time(self.scheduled_table['start time (UTC)'][i])).value * 24*60

                    end_scheduled_table = end_scheduled_table.append({'target': self.scheduled_table['target'][i]+'_2',
                                                                      'start time (UTC)':
                                                                          self.SS1_night_blocks['end time (UTC)'][0],
                                                                      'end time (UTC)': end_before_cut,
                                                                      'duration (minutes)': (Time(end_before_cut)-Time(
                                                                          self.SS1_night_blocks['end time (UTC)'][
                                                                              0])).value * 24 * 60,
                                                                      'ra (h)': self.scheduled_table['ra (h)'][i],
                                                                      'ra (m)': self.scheduled_table['ra (m)'][i],
                                                                      'ra (s)': self.scheduled_table['ra (s)'][i],
                                                                      'dec (d)': self.scheduled_table['dec (d)'][i],
                                                                      'dec (m)': self.scheduled_table['dec (m)'][i],
                                                                      'dec (s)': self.scheduled_table['dec (s)'][i],
                                                                      'configuration':
                                                                          self.scheduled_table['configuration'][i]},
                                                                     ignore_index=True)

                if self.SS1_night_blocks['end time (UTC)'][0] >= end_before_cut:

                    # situation 7
                    if (self.SS1_night_blocks['end time (UTC)'][0] >= 
                            self.end_of_observation.iso):
                        
                        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 7')

                        if self.SS1_night_blocks['start time (UTC)'][0] > self.scheduled_table['end time (UTC)'][i]:
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK +
                                  ' situation 7a, no change made to initial schedule')
                        else:
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 7b')
                            self.scheduled_table['end time (UTC)'][i] = self.SS1_night_blocks['start time (UTC)'][0]
                            self.scheduled_table['duration (minutes)'] = \
                                (Time(self.scheduled_table['end time (UTC)'][i])
                                 - Time(self.scheduled_table['start time (UTC)'][i])).value * 24 * 60

                            self.SS1_night_blocks['end time (UTC)'][0] = \
                                self.end_of_observation.iso
                            self.SS1_night_blocks['duration (minutes)'][0] = \
                                (Time(self.SS1_night_blocks['end time (UTC)'][0]) -
                                 Time(self.SS1_night_blocks['start time (UTC)'][0])).value * 24*60

                    elif (self.SS1_night_blocks['end time (UTC)'][0] <=
                          self.end_of_observation.iso):
                        
                        # situation 8
                        if self.SS1_night_blocks['start time (UTC)'][0] >= end_before_cut:
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK +
                                  ' situation 8, no change made to initial schedule')
                        # situation 9
                        else:
                            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 9')
                            self.scheduled_table['end time (UTC)'][i] = self.SS1_night_blocks['start time (UTC)'][0]
                            self.scheduled_table['duration (minutes)'] = \
                                (Time(self.scheduled_table['end time (UTC)'][i]) -
                                 Time(self.scheduled_table['start time (UTC)'][i])).value * 24 * 60

                # situation 10
                if self.SS1_night_blocks['start time (UTC)'][0] >= end_before_cut:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' situation 10, no change made to initial schedule')

            end_scheduled_table = end_scheduled_table.append({'target': self.SS1_night_blocks['target'][0],
                                                              'start time (UTC)': 
                                                                  self.SS1_night_blocks['start time (UTC)'][0],
                                                              'end time (UTC)': 
                                                                  self.SS1_night_blocks['end time (UTC)'][0],
                                                              'duration (minutes)': 
                                                                  self.SS1_night_blocks['duration (minutes)'][0],
                                                              'ra (h)': self.SS1_night_blocks['ra (h)'][0],
                                                              'ra (m)': self.SS1_night_blocks['ra (m)'][0],
                                                              'ra (s)': self.SS1_night_blocks['ra (s)'][0],
                                                              'dec (d)': self.SS1_night_blocks['dec (d)'][0],
                                                              'dec (m)': self.SS1_night_blocks['dec (m)'][0],
                                                              'dec (s)': self.SS1_night_blocks['dec (s)'][0],
                                                              'configuration': 
                                                                  self.SS1_night_blocks['configuration'][0]},
                                                             ignore_index=True)
            end_scheduled_table = end_scheduled_table.append({'target': self.scheduled_table['target'][i],
                                                              'start time (UTC)': 
                                                                  self.scheduled_table['start time (UTC)'][i],
                                                              'end time (UTC)': 
                                                                  self.scheduled_table['end time (UTC)'][i],
                                                              'duration (minutes)': 
                                                                  self.scheduled_table['duration (minutes)'][i],
                                                              'ra (h)': self.scheduled_table['ra (h)'][i],
                                                              'ra (m)': self.scheduled_table['ra (m)'][i],
                                                              'ra (s)': self.scheduled_table['ra (s)'][i],
                                                              'dec (d)': self.scheduled_table['dec (d)'][i],
                                                              'dec (m)': self.scheduled_table['dec (m)'][i],
                                                              'dec (s)': self.scheduled_table['dec (s)'][i],
                                                              'configuration': 
                                                                  self.scheduled_table['configuration'][i]},
                                                             ignore_index=True)

        end_scheduled_table = Table.from_pandas(end_scheduled_table)
        end_scheduled_table = unique(end_scheduled_table, keys='target')
        end_scheduled_table.sort(keys='start time (UTC)')
        idx_too_short_block = np.where((end_scheduled_table['duration (minutes)'] <= 3))

        if idx_too_short_block:
            for i in list(idx_too_short_block[0]):
                idx_for_target_2 = \
                    np.where((end_scheduled_table['target'] == end_scheduled_table['target'][i] + '_2'))[0]
                if idx_for_target_2.size != 0:
                    end_scheduled_table['target'][idx_for_target_2[0]] = end_scheduled_table['target'][i]

            list_too_short_reverse = list(idx_too_short_block[0])
            list_too_short_reverse.reverse()  # important to reverse if you want to delete lines correctly
            for i in list_too_short_reverse:
                end_scheduled_table.remove_row(i)

        self.scheduled_table_sorted = end_scheduled_table
        # ***** Update target lists on server *****
        # subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/target_lists/target_list_special.txt',
        #                   'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/target_lists/'])
        # subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/target_lists/target_transit_follow_up.txt',
        #                   'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/target_lists/'])

        return self.scheduled_table_sorted

    def make_scheduled_table(self):
        path = path_spock + '/DATABASE'
        try:
            os.path.exists(os.path.join(path, self.telescope, 'night_blocks_' + str(self.telescope) + '_' +
                                        self.day_of_night.tt.datetime[0].strftime("%Y-%m-%d") + '.txt'))
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Local path exists and is: ',
                  os.path.join(path, self.telescope, 'night_blocks_' + str(self.telescope) + '_' +
                               self.day_of_night.tt.datetime[0].strftime("%Y-%m-%d") + '.txt'))
        except TypeError:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Local path does not exist yet ')
        except NameError:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no input night_block for this day')
        except FileNotFoundError:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no input night_block for this day')

        if not (self.scheduled_table is None):
            return self.scheduled_table
        else:
            try:
                self.scheduled_table = read_night_block(telescope=self.telescope,
                                                        day=self.day_of_night.tt.datetime.strftime("%Y-%m-%d"))
                return self.scheduled_table
            except TypeError:
                self.scheduled_table = read_night_block(telescope=self.telescope,
                                                        day=self.day_of_night.tt.datetime.strftime("%Y-%m-%d"))
                return self.scheduled_table

    def make_night_block(self):

        path = path_spock + '/night_blocks_propositions/'
        if not (self.scheduled_table_sorted is None):
            if isinstance(self.scheduled_table_sorted, pd.DataFrame):
                self.scheduled_table_sorted = self.scheduled_table_sorted.set_index('target')
                try:
                    self.scheduled_table_sorted.drop('TransitionBlock', inplace=True)
                except KeyError:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transition block')
                panda_table = self.scheduled_table_sorted
                try:
                    panda_table.to_csv(os.path.join(path, 'night_blocks_' + str(self.telescope) + '_' +
                                                    self.day_of_night.tt.datetime[0].strftime("%Y-%m-%d") +
                                                    '.txt'), sep=' ')
                except TypeError:
                    panda_table.to_csv(os.path.join(path, 'night_blocks_' + str(self.telescope) + '_' +
                                                    self.day_of_night.tt.datetime.strftime("%Y-%m-%d") +
                                                    '.txt'), sep=' ')
            else:
                try:
                    self.scheduled_table_sorted.add_index('target')
                    index_to_delete = self.scheduled_table_sorted.loc['TransitionBlock'].index
                    self.scheduled_table_sorted.remove_row(index_to_delete)
                except KeyError:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transition block')
                panda_table = self.scheduled_table_sorted.to_pandas()
                try:
                    panda_table.to_csv(os.path.join(path, 'night_blocks_' + str(self.telescope) + '_' +
                                                    self.day_of_night.tt.datetime[0].strftime("%Y-%m-%d") +
                                                    '.txt'), sep=' ')
                except TypeError:
                    panda_table.to_csv(os.path.join(path, 'night_blocks_' + str(self.telescope) + '_' +
                                                    self.day_of_night.tt.datetime.strftime("%Y-%m-%d") +
                                                    '.txt'), sep=' ')

    def exposure_time(self, input_name, target_list, day=None):
        """

        Parameters
        ----------
        input_name
        target_list
        day

        Returns
        -------

        """
        i = np.where((target_list['Sp_ID'] == input_name))[0]
        if day is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
        # moon_phase = round(moon_illumination(Time(day.iso, out_subfmt='date')), 2)
        try:
            spectral_type = round(float(target_list['SpT'][i].data.data[0]))
        except AttributeError:
            try:
                spectral_type = round(float(target_list['SpT'][i]))
            except ValueError:
                spectral_type = target_list['SpT'][i].values[0]
        except NotImplementedError:
            try:
                spectral_type = round(float(target_list['SpT'][i]))
            except ValueError:
                spectral_type = target_list['SpT'][i].values[0]
        if not isinstance(spectral_type, str):
            if round(float(abs(target_list['SpT'][i]))) <= 9:
                spt_type = 'M' + str(round(float(abs(target_list['SpT'][i]))))
                if spt_type == 'M3':
                    spt_type = 'M2'
            if round(float(abs(target_list['SpT'][i]))) <= 2:
                spt_type = 'M2'
            elif (round(float(abs(target_list['SpT'][i]))) == 12) or (round(float(abs(target_list['SpT'][i]))) == 15)\
                    or (round(float(abs(target_list['SpT'][i])) == 18)):
                spt_type = 'M' + str(round(target_list['SpT'][i])) - 10
            elif round(float(abs(target_list['SpT'][i]))) == 10:
                spt_type = 'M9'
            elif round(float(abs(target_list['SpT'][i]))) == 11:
                spt_type = 'L2'
            elif round(float(abs(target_list['SpT'][i]))) == 13:
                spt_type = 'L2'
            elif round(float(abs(target_list['SpT'][i]))) == 14:
                spt_type = 'L5'
            elif round(float(abs(target_list['SpT'][i]))) > 14:
                spt_type = 'L8'
        else:
            spt_type = spectral_type

        filt_ = target_list['Filter_spc'][i].values[0]
        if (filt_ == 'z\'') or (filt_ == 'r\'') or (filt_ == 'i\'') or (filt_ == 'g\''):
            filt_ = filt_.replace('\'', '')
        if filt_ != 'I+z':
            filters = [filt_] + ['I+z', 'z', 'i', 'r']
        else:
            filters = ['I+z', 'z', 'i', 'r']
        filt_idx = 0
        filt_ = filters[filt_idx]
        texp = 0

        while texp < 10:

            if 0 < filt_idx <= 3:
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Change filter to avoid saturation!!')
                filt_ = filters[filt_idx]
                print(filt_)

            if self.telescope == 'Saint-Ex':
                print(self.day_of_night)
                moon_phase = round(moon_illumination(Time(self.day_of_night.iso, out_subfmt='date')), 2)
                if float(target_list['J'][i]) != 0.:
                    a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
                                 filt=filt_, airmass=1.1, moonphase=moon_phase, irtf=0.8, num_tel=1,
                                 seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
                else:
                    if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
                        a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
                                     filt=filt_, airmass=1.1, moonphase=moon_phase, irtf=0.8, num_tel=1,
                                     seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
                    else:
                        sys.exit('ERROR: You must precise Vmag or Jmag for this target')
                texp = a.exp_time_calculator(ADUpeak=30000)[0]

            elif self.telescope == 'TN_Oukaimeden':
                if float(target_list['J'][i]) != 0.:
                    a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
                                 filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
                else:
                    if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
                        a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
                                     filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
                    else:
                        sys.exit('ERROR: You must precise Vmag or Jmag for this target')
                texp = a.exp_time_calculator(ADUpeak=50000)[0]
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
                      ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')

            elif self.telescope == 'TS_La_Silla':
                if float(target_list['J'][i]) != 0.:
                    a = (ETC.etc(mag_val=target_list['J'][i], mag_band='J', spt=spt_type, filt=filt_,
                                 airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.4, gain=1.1))
                else:
                    if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
                        a = (ETC.etc(mag_val=target_list['V'][i], mag_band='V', spt=spt_type, filt=filt_,
                                     airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.4, gain=1.1))
                    else:
                        sys.exit('ERROR: You must precise Vmag or Jmag for ' + str(target_list['Sp_ID'][i]))
                texp = a.exp_time_calculator(ADUpeak=50000)[0]
                print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Don\'t forget to  '
                                                               'calculate exposure time for TRAPPIST observations!!')

            elif self.telescope == 'Artemis':
                if float(target_list['J'][i]) != 0.:
                    a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
                                 filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
                else:
                    if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
                        a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
                                     filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
                    else:
                        sys.exit('ERROR: You must precise Vmag or Jmag for this target')
                texp = a.exp_time_calculator(ADUpeak=45000)[0]

            elif self.telescope == 'Io' or self.telescope == 'Europa' \
                    or self.telescope == 'Ganymede' or self.telescope == 'Callisto':
                if float(target_list['J'][i]) != 0.:
                    a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
                                 filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
                else:
                    if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
                        a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
                                     filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
                    else:
                        sys.exit('ERROR: You must precise Vmag or Jmag for this target')
                texp = a.exp_time_calculator(ADUpeak=45000)[0]

            if filt_idx > 3:
                print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' You have to defocus in we want to observe this target')
                texp = 10.0001
                filt_ = 'r'

            filt_idx += 1

        target_list['Filter_spc'][i] = filt_

        return int(texp), filt_

    # def exposure_time_st(self, input_name, target_list, day=None):
    #     i = np.where((target_list['Sp_ID'] == input_name))[0]
    #     spt_type = ''
    #     texp = 10
    #     if day is None:
    #         print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
    #     # moon_phase = round(moon_illumination(Time(day.iso, out_subfmt='date')), 2)
    #     try:
    #         spectral_type = round(float(target_list['SpT'][i].data.data[0]))
    #     except AttributeError:
    #         try:
    #             spectral_type = round(float(target_list['SpT'][i]))
    #         except ValueError:
    #             spectral_type = target_list['SpT'][i].values[0]
    #     except NotImplementedError:
    #         try:
    #             spectral_type = round(float(target_list['SpT'][i]))
    #         except ValueError:
    #             spectral_type = target_list['SpT'][i].values[0]
    #     if not isinstance(spectral_type, str):
    #         if spectral_type <= 9:
    #             spt_type = 'M' + str(spectral_type)
    #             if spt_type == 'M3':
    #                 spt_type = 'M2'
    #         elif (spectral_type == 12) or (spectral_type == 15) or (
    #                 int(spectral_type) == 18):
    #             spt_type = 'M' + str(spectral_type - 10)
    #         elif spectral_type == 10:
    #             spt_type = 'M9'
    #         elif spectral_type == 11:
    #             spt_type = 'L2'
    #         elif spectral_type == 13:
    #             spt_type = 'L2'
    #         elif spectral_type == 14:
    #             spt_type = 'L5'
    #         elif spectral_type > 14:
    #             spt_type = 'L8'
    #     else:
    #         spt_type = spectral_type
    #     filt_ = target_list['Filter_spc'][i].values[0]
    #     if (filt_ == 'z\'') or (filt_ == 'r\'') or (filt_ == 'i\'') or (filt_ == 'g\''):
    #         filt_ = filt_.replace('\'', '')
    #     if filt_ != 'I+z':
    #         filters = [filt_] + ['I+z', 'z', 'i', 'r', 'g']
    #     else:
    #         filters = ['I+z', 'z', 'i', 'r', 'g']
    #     filt_idx = 0
    #     filt_ = filters[filt_idx]
    #     if self.telescope == 'Saint-Ex':
    #         if float(target_list['J'][i]) != 0.:
    #             a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                          filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1,
    #                          seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
    #         else:
    #             if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                 a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1,
    #                              seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
    #             else:
    #                 sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #         texp = a.exp_time_calculator(ADUpeak=30000)[0]
    #
    #     elif self.telescope == 'TS_La_Silla' or self.telescope == 'TN_Oukaimeden':
    #         if float(target_list['J'][i]) != 0.:
    #             a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                          filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #         else:
    #             if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                 a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #             else:
    #                 sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #         texp = a.exp_time_calculator(ADUpeak=50000)[0]
    #         print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #               ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')
    #
    #     elif self.telescope == 'Artemis':
    #         if float(target_list['J'][i]) != 0.:
    #             a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                          filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #         else:
    #             if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                 a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #             else:
    #                 sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #         texp = a.exp_time_calculator(ADUpeak=45000)[0]
    #
    #     elif self.telescope == 'Io' or self.telescope == 'Europa' \
    #             or self.telescope == 'Ganymede' or self.telescope == 'Callisto':
    #         if float(target_list['J'][i]) != 0.:
    #             a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                          filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
    #         else:
    #             if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                 a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
    #             else:
    #                 sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #         texp = a.exp_time_calculator(ADUpeak=45000)[0]
    #
    #     while texp < 10:
    #         print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' Change filter to avoid saturation!!')
    #         filt_idx += 1
    #         if filt_idx >= 3:
    #             print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' You have to defocus in we want to observe this target')
    #             texp = 10.0001
    #         filt_ = filters[filt_idx]
    #         if self.telescope == 'Saint-Ex':
    #             if float(target_list['J'][i]) != 0.:
    #                 a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1,
    #                              seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
    #             else:
    #                 if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                     a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
    #                                  filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1,
    #                                  seeing=0.7, gain=3.48, temp_ccd=-70, observatory_altitude=2780))
    #                 else:
    #                     sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #             texp = a.exp_time_calculator(ADUpeak=30000)[0]
    #
    #         elif self.telescope == 'Artemis':
    #             target_list['Filter_spc'][i] = filt_
    #             if float(target_list['J'][i]) != 0.:
    #                 a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #             else:
    #                 if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                     a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type,
    #                                  filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #                 else:
    #                     sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #             texp = a.exp_time_calculator(ADUpeak=45000)[0]
    #
    #         elif self.telescope == 'TS_La_Silla' or self.telescope == 'TN_Oukaimeden':
    #             target_list['Filter_spc'][i] = filt_
    #             if float(target_list['J'][i]) != 0.:
    #                 a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #             else:
    #                 if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                     a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                                  filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=1.0, gain=1.1))
    #                 else:
    #                     sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #             texp = a.exp_time_calculator(ADUpeak=50000)[0]
    #             print(
    #                 Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #                 ' Don\'t forget to  calculate exposure time for TRAPPIST observations!!')
    #
    #         elif self.telescope == 'Io' or self.telescope == 'Europa' or \
    #                 self.telescope == 'Ganymede' or self.telescope == 'Callisto':
    #
    #             target_list['Filter_spc'][i] = filt_
    #             if float(target_list['J'][i]) != 0.:
    #                 a = (ETC.etc(mag_val=float(target_list['J'][i]), mag_band='J', spt=spt_type,
    #                              filt=filt_, airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
    #             else:
    #                 if (float(target_list['J'][i]) == 0.) and (float(target_list['V'][i]) != 0.):
    #                     a = (ETC.etc(mag_val=float(target_list['V'][i]), mag_band='V', spt=spt_type, filt=filt_,
    #                                  airmass=1.1, moonphase=0.5, irtf=0.8, num_tel=1, seeing=0.7, gain=1.1))
    #                 else:
    #                     sys.exit('ERROR: You must precise Vmag or Jmag for this target')
    #             texp = a.exp_time_calculator(ADUpeak=45000)[0]
    #
    #         if filt_idx > 3:
    #             print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' You have to defocus in we want to observe this target')
    #             texp = 10.0001
    #             filt_ = 'r'
    #
    #     target_list['Filter_spc'][i] = filt_
    #
    #     return int(texp), filt_

    def visibility_plot(self):
        day = self.day_of_night
        delta_midnight = Time(np.linspace(self.start_of_observation.jd,
                                          self.end_of_observation.jd, 100),
                              format='jd')
        for i in range(len(self.scheduled_table_sorted)):
            idx = np.where((self.target_table_spc['Sp_ID'] == self.scheduled_table_sorted['target'][i]))[0]
            plot_airmass(self.targets[int(idx)], self.observatory, delta_midnight, style_sheet=dark_style_sheet)
            t = Time(self.scheduled_table_sorted['start time (UTC)'][i])
            plt.vlines(t.iso, 3, 1, color='r')
        plt.legend(shadow=True, loc=2)


def visibility_plot(day, observatory, night_block):
    delta_midnight = Time(np.linspace(observatory.twilight_evening_nautical(day, which='next').jd,
                                      observatory.twilight_morning_nautical(day+1, which='nearest').jd, 100),
                          format='jd')
    dec = str(int(float(night_block['dec (d)'][0]))) + 'd' + str(int(float(night_block['dec (m)'][0]))) + 'm' + \
          str(int(float(night_block['dec (s)'][0]))) + 's'
    ra = str(int(float(night_block['ra (h)'][0]))) + 'h' + str(int(float(night_block['ra (m)'][0]))) + 'm' + \
         str(int(float(night_block['ra (s)'][0]))) + 's'
    for i in range(len(night_block)):
        plot_airmass(SkyCoord(ra=ra, dec=dec, frame='icrs'), observatory, delta_midnight, style_sheet=dark_style_sheet)
        t = Time(night_block['start time (UTC)'][i])
        plt.vlines(t.iso, 3, 1, linestyle='-', color='r', alpha=0.7)
        plt.legend(shadow=True, loc=2)


def save_schedule(save, over_write, day, telescope):
    if save:
        source = path_spock + '/night_blocks_propositions/' + 'night_blocks_' + telescope + '_' +\
                 day.tt.datetime.strftime("%Y-%m-%d") + '.txt'
        destination = path_spock + '/DATABASE/' + telescope + '/'
        destination_2 = path_spock + '/DATABASE/' + telescope + '/' + 'Archive_night_blocks/'
        if over_write:
            dest = shutil.copy(source, destination)
            dest2 = shutil.copy(source, destination_2)
            print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been over-written to ' +
                  '\"' + destination + '\"')
        if not over_write:
            try:
                dest = shutil.move(source, destination)
                dest2 = shutil.move(source, destination_2)
                print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been copied to ' +
                      '\"' + destination + '\"')
            except shutil.Error:
                print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + destination + 'night_blocks_' +
                      telescope + '_' + day.tt.datetime.strftime("%Y-%m-%d") + '.txt' + '\"' + ' already exists')
    if not save:
        print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + ' Those plans have not been saved')


def make_plans(day, nb_days, telescope):
    make_np(day, nb_days, telescope)


def upload_plans(day, nb_days, telescope):
    if telescope.find('Callisto') is not -1:
        upload_np_calli(day, nb_days)
    if telescope.find('Ganymede') is not -1:
        upload_np_gany(day, nb_days)
    if telescope.find('Io') is not -1:
        upload_np_io(day, nb_days)
    if telescope.find('Europa') is not -1:
        upload_np_euro(day, nb_days)
    if telescope.find('Artemis') is not -1:
        upload_np_artemis(day, nb_days)
    if telescope.find('TS_La_Silla') is not -1:
        upload_np_ts(day, nb_days)
    if telescope.find('TN_Oukaimeden') is not -1:
        upload_np_tn(day, nb_days)
    # ------------------- update archive date by date plans folder  ------------------

    path_gant_chart = os.path.join(path_spock + '/SPOCK_Figures/Preview_schedule.html')
    path_database_home = \
        os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/Preview_schedule.html')
    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path local \'Gant chart\' = ', path_gant_chart)
    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path database = \'Gant chart\' = ',  path_database_home)
    subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart, path_database_home])
    path_gant_chart_masterfile = \
        os.path.join('/Users/elsaducrot/spock_2/SPOCK_Figures/spock_stats_masterfile.csv')
    path_database_home_masterfile = \
        os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_stats_masterfile.csv')
    subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart_masterfile,
                      path_database_home_masterfile])


def read_night_block(telescope, day):
    day_fmt = Time(day, scale='utc', out_subfmt='date').tt.datetime.strftime("%Y-%m-%d")
    path_local = path_spock + '/DATABASE/' + telescope+'/Archive_night_blocks/night_blocks_' + telescope + '_' + \
                 day_fmt + '.txt'

    if os.path.exists(path_local):
        day_fmt = Time(day, scale='utc', out_subfmt='date').tt.datetime.strftime("%Y-%m-%d")
        scheduler_table = Table.read(
            path_spock + '/DATABASE/' + str(telescope) + '/Archive_night_blocks' + '/night_blocks_' +
            str(telescope) + '_' + str(day_fmt) + '.txt',
            format='ascii')
    else:
        nightb_url = "http://www.mrao.cam.ac.uk/SPECULOOS/"+telescope+'/schedule/Archive_night_blocks/night_blocks_' + \
                     telescope+'_'+day_fmt+'.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            open(path_local, 'wb').write(nightb.content)
            scheduler_table = pd.read_csv(path_local, delimiter=' ',
                                          skipinitialspace=True, error_bad_lines=False)
    return scheduler_table


def date_range_in_days(date_range):
    date_format = "%Y-%m-%d %H:%M:%S.%f"
    date_start = datetime.strptime(date_range[0].value, date_format)
    date_end = datetime.strptime(date_range[1].value, date_format)
    return (date_end - date_start).days


def get_info_follow_up_target(name, target_list_follow_up):
    """

    Parameters
    ----------
    name: name of the target  you wish to the next transit windows of
    target_list_follow_up: dataframe

    Returns
    -------

    """
    idx_target = np.where((target_list_follow_up['Sp_ID'] == name))[0]
    if idx_target.size == 0:
        sys.exit('ERROR: This target name is not in the list.')
    else:
        return float(target_list_follow_up['RA'][idx_target]), float(target_list_follow_up['DEC'][idx_target]),\
               float(target_list_follow_up['T0'][idx_target]), float(target_list_follow_up['P'][idx_target]),\
               float(target_list_follow_up['W'][idx_target])


def prediction(name, ra, dec, timing, period, duration, start_date, ntr):
    altitude_constraint = 25
    altitude_constraint_stx = 28

    start_date = Time(start_date)

    constraints_predictions = [AltitudeConstraint(min=altitude_constraint * u.deg),
                               AtNightConstraint.twilight_nautical()]
    constraints_predictions_stx = [AltitudeConstraint(min=altitude_constraint_stx * u.deg),
                               AtNightConstraint.twilight_nautical()]

    target_transit = EclipsingSystem(primary_eclipse_time=Time(timing, format='jd'),
                                     orbital_period=period*u.day, duration=duration*u.day,
                                     name=name)
    target = [FixedTarget(coord=SkyCoord(ra=ra* u.degree, dec=dec* u.degree), name=name)]


    target_set_time_sso = charge_observatories('SSO')[0].target_set_time(start_date,
                                                                         target,
                                                                         which='next',
                                                                         horizon=altitude_constraint * u.deg)


    mid_transit_timing = Time(target_transit.next_primary_eclipse_time(start_date, n_eclipses=ntr)).iso
    mid_transit_timing_jd = Time(target_transit.next_primary_eclipse_time(start_date, n_eclipses=ntr)).jd

    ing_egr = target_transit.next_primary_ingress_egress_time(start_date, n_eclipses=ntr)

    target = [FixedTarget(coord=SkyCoord(ra=ra * u.degree, dec=dec * u.degree), name=name)]
    # observatory = charge_observatories(obs_name)[0]

    observable_sso = is_event_observable(constraints_predictions, charge_observatories('SSO')[0], target,
                                         times_ingress_egress=ing_egr)
    observable_sno = is_event_observable(constraints_predictions, charge_observatories('SNO')[0], target,
                                         times_ingress_egress=ing_egr)
    observable_saintex = is_event_observable(constraints_predictions_stx, charge_observatories('Saint-Ex')[0], target,
                                             times_ingress_egress=ing_egr)
    observable_tn_oukaimeden = is_event_observable(constraints_predictions, charge_observatories('TN_Oukaimeden')[0],
                                                   target, times_ingress_egress=ing_egr)
    observable_ts_la_silla = is_event_observable(constraints_predictions, charge_observatories('TS_La_Silla')[0],
                                                 target, times_ingress_egress=ing_egr)

    dates = [Time(Time(ing_egr[i][0]).iso, out_subfmt='date').iso for i in range(len(ing_egr))]

    observable_sso_tables = []
    observable_sno_tables = []
    observable_saint_ex_tables = []

    for i in range(len(ing_egr)):
        date = Time(Time(Time(Time(ing_egr[i][0]).iso, out_subfmt='date').iso).jd+0.625, format='jd')
        observable_sso_tables.append(observability_table(constraints_predictions, charge_observatories('SSO')[0],
                                                         target, time_range=Time([date, date+1]),
                                                         time_grid_resolution=0.5 * u.hour))

        observable_sno_tables.append(observability_table(constraints_predictions, charge_observatories('SNO')[0],
                                                         target, time_range=Time([date, date + 1]),
                                                         time_grid_resolution=0.5 * u.hour))

        observable_saint_ex_tables.append(observability_table(constraints_predictions_stx,
                                                              charge_observatories('Saint-Ex')[0],
                                                              target, time_range=Time([date, date + 1]),
                                                              time_grid_resolution=0.5 * u.hour))

    observable_sso_table = vstack(observable_sso_tables)
    observable_sno_table = vstack(observable_sno_tables)
    observable_saint_ex_table = vstack(observable_saint_ex_tables)

    df = pd.DataFrame(data=ing_egr.sort().iso, index=[name]*len(ing_egr), columns=["Ingress", "Egress"])

    df['mid transit'] = mid_transit_timing

    df['mid transit JD'] = np.round(mid_transit_timing_jd, 3)

    df['Observable SSO'] = observable_sso[0]

    df['Hours observable SSO (h)'] = observable_sso_table['fraction of time observable'] * 24*u.hour

    df['Observable SNO'] = observable_sno[0]

    df['Hours observable SNO (h)'] = observable_sno_table['fraction of time observable'] * 24*u.hour

    df['Observable Saint-Ex'] = observable_saintex[0]

    df['Hours observable Saint-Ex (h)'] = observable_saint_ex_table['fraction of time observable'] * 24*u.hour

    df['Observable TN Oukaimeden'] = observable_tn_oukaimeden[0]

    df['Observable TS La Silla'] = observable_ts_la_silla[0]

    return df


# def make_docx_schedule(observatory, telescope, date_range, name_operator):
#
#     if not os.path.exists(path_spock + '/TRAPPIST_schedules_docx'):
#         os.makedirs(path_spock + '/TRAPPIST_schedules_docx')
#
#     df_speculoos = pd.read_csv(target_list_from_stargate_path, delimiter=',')
#     df_follow_up = pd.read_csv(path_spock + '/target_lists/target_transit_follow_up.txt', delimiter=' ')
#     df_special = pd.read_csv(path_spock + '/target_lists/target_list_special.txt', delimiter=' ')
#
#     df_follow_up['nb_hours_surved'] = [0]*len(df_follow_up)
#     df_follow_up['nb_hours_threshold'] = [0] * len(df_follow_up)
#     df_special['nb_hours_surved'] = [0] * len(df_special)
#     df_special['nb_hours_threshold'] = [0] * len(df_special)
#
#     df_pandas = pd.DataFrame({'Sp_ID': pd.concat([df_speculoos['Sp_ID'], df_follow_up['Sp_ID'], df_special['Sp_ID']]),
#                               'RA': pd.concat([df_speculoos['RA'], df_follow_up['RA'], df_special['RA']]),
#                               'DEC': pd.concat([df_speculoos['DEC'], df_follow_up['DEC'], df_special['DEC']]),
#                               'J': pd.concat([df_speculoos['J'], df_follow_up['J'], df_special['J']]),
#                               'SpT': pd.concat([df_speculoos['SpT'], df_follow_up['SpT'], df_special['SpT']]),
#                               'nb_hours_surved': pd.concat([df_speculoos['nb_hours_surved'],
#                                                             df_follow_up['nb_hours_surved'],
#                                                             df_special['nb_hours_surved']]),
#                               'nb_hours_threshold': pd.concat([df_speculoos['nb_hours_threshold'],
#                                                                df_follow_up['nb_hours_threshold'],
#                                                                df_special['nb_hours_threshold']])})
#     df_pandas = df_pandas.drop_duplicates()
#     df = Table.from_pandas(df_pandas)
#     nb_day_date_range = date_range_in_days(date_range)
#     doc = Document()
#     par = doc.add_paragraph()
#     par_format = par.paragraph_format
#     par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     par_format.space_beFore = Pt(0)
#     par_format.space_after = Pt(6)
#     run = par.add_run(observatory.name)
#     run.bold = True
#     font = run.font
#     font.size = Pt(16)
#     font.color.rgb = RGBColor(0, 0, 0)
#     par = doc.add_paragraph()
#     par_format = par.paragraph_format
#     par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     par_format.space_beFore = Pt(0)
#     par_format.space_after = Pt(12)
#     run = par.add_run('Schedule from ' + Time(date_range[0].iso, out_subfmt='date').iso + ' to ' +
#                       Time(date_range[1].iso, out_subfmt='date').iso)
#     run.bold = True
#     font = run.font
#     font.size = Pt(16)
#     font.color.rgb = RGBColor(0, 0, 0)
#     par = doc.add_paragraph()
#     par_format = par.paragraph_format
#     par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     par_format.space_beFore = Pt(0)
#     par_format.space_after = Pt(12)
#     run = par.add_run('(Total time = 0hr, technical loss = 0hr, weather loss = 0hr,'
#                       'Exotime = 0hr, cometime = 0hr, chilean time = 0hr)')
#     run.bold = True
#     font = run.font
#     font.size = Pt(12)
#     font.color.rgb = RGBColor(255, 0, 0)
#     par = doc.add_paragraph()
#     par_format = par.paragraph_format
#     par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     par_format.space_beFore = Pt(0)
#     par_format.space_after = Pt(20)
#     run = par.add_run(name_operator)
#     run.italic = True
#     font = run.font
#     font.size = Pt(12)
#     font.color.rgb = RGBColor(0, 0, 0)
#     par = doc.add_paragraph()
#     par_format = par.paragraph_format
#     par_format.space_beFore = Pt(16)
#     par_format.space_after = Pt(0)
#
#     for i in range(nb_day_date_range):
#
#         date = date_range[0] + i
#         table_schedule = read_night_block(telescope, date)
#         sun_set = observatory.sun_set_time(date, which='next').iso
#         sun_rise = observatory.sun_rise_time(date, which='next').iso
#         moon_illum = int(round(moon_illumination(date) * 100, 0)) * u.percent
#         civil_twilights = [Time(observatory.twilight_evening_civil(date, which='next')).iso,
#                            Time(observatory.twilight_morning_civil(date + 1, which='nearest')).iso]
#         nautic_twilights = [Time(observatory.twilight_evening_nautical(date, which='next')).iso,
#                             Time(observatory.twilight_morning_nautical(date + 1, which='nearest')).iso]
#         astro_twilights = [Time(observatory.twilight_evening_astronomical(date, which='next')).iso,
#                            Time(observatory.twilight_morning_astronomical(date + 1, which='nearest')).iso]
#         start_night = table_schedule['start time (UTC)'][0]
#         end_night = np.array(table_schedule['end time (UTC)'])[-1]
#         night_duration = round((Time(nautic_twilights[1]) - Time(nautic_twilights[0])).jd * 24, 3) * u.hour
#
#         run = par.add_run('Night starting on the ' + Time(date, out_subfmt='date').value)
#         run.bold = True
#         run.underline = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#         par = doc.add_paragraph()
#         par_format = par.paragraph_format
#         par_format.space_beFore = Pt(0)
#         par_format.space_after = Pt(0)
#         run = par.add_run('Moon illumination: ' + str(moon_illum))
#         run.italic = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#         par = doc.add_paragraph()
#         par_format = par.paragraph_format
#         par_format.space_beFore = Pt(0)
#         par_format.space_after = Pt(0)
#         run = par.add_run(
#             'Sunset - Sunrise: ' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.hour) +
#             'h' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.minute) +
#             '  / ' + '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.hour) + 'h' +
#             '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.minute))
#         run.italic = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#         par = doc.add_paragraph()
#         par_format = par.paragraph_format
#         par_format.space_beFore = Pt(0)
#         par_format.space_after = Pt(0)
#         run = par.add_run(
#             'Civil/Naut./Astro. twilights: ' +
#             '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' +
#             '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.minute) +
#             '-' + '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.hour) +
#             'h' + '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.minute) +
#             ' / ' + '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.hour) +
#             'h' + '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.minute) +
#             '-' + '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.hour) +
#             'h' + '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.minute) +
#             '  / ' + '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.hour) +
#             'h' + '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.minute) +
#             '-' + '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' +
#             '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.minute))
#         run.italic = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#         par = doc.add_paragraph()
#         par_format = par.paragraph_format
#         par_format.space_beFore = Pt(0)
#         par_format.space_after = Pt(0)
#         run = par.add_run('Start-end of night (Naut. twil.): ' + '{:02d}'.format(Time(start_night).datetime.hour) +
#                           'h' + '{:02d}'.format(Time(start_night).datetime.minute) +
#                           ' to ' + '{:02d}'.format(Time(end_night).datetime.hour) + 'h' +
#                           '{:02d}'.format(Time(end_night).datetime.minute))
#         run.italic = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#         par = doc.add_paragraph()
#         par_format = par.paragraph_format
#         par_format.space_beFore = Pt(0)
#         par_format.space_after = Pt(3)
#         run = par.add_run('Night duration (Naut. twil.): ' + str(night_duration))
#         run.italic = True
#         font = run.font
#         font.size = Pt(12)
#         font.color.rgb = RGBColor(0, 0, 0)
#
#         for j in range(len(table_schedule)):
#             trappist_planets = ['Trappist-1b', 'Trappist-1c', 'Trappist-1d', 'Trappist-1e',
#                                 'Trappist-1f', 'Trappist-1g', 'Trappist-1h']
#
#             if any(table_schedule['target'][j] == p for p in trappist_planets):
#                 idx_target = np.where((df['Sp_ID'] == 'Sp2306-0502'))[0]
#             else:
#                 idx_target = np.where((df['Sp_ID'] == table_schedule['target'][j]))[0]
#
#             start_time_target = table_schedule['start time (UTC)'][j]
#             end_time_target = table_schedule['end time (UTC)'][j]
#             config = table_schedule['configuration'][j]
#             try:
#                 coords = SkyCoord(ra=df['RA'][idx_target].data.data[0] * u.deg,
#                                   dec=df['DEC'][idx_target].data.data[0] * u.deg)
#             except IndexError:
#                 break
#             dist_moon = '34'
#
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(0)
#             par_format.space_after = Pt(0)
#             run = par.add_run(
#                 'From ' + '{:02d}'.format(Time(start_time_target, out_subfmt='date_hm').datetime.hour) +
#                 'h' + '{:02d}'.format(Time(start_time_target, out_subfmt='date_hm').datetime.minute) +
#                 ' to ' + '{:02d}'.format(Time(end_time_target, out_subfmt='date_hm').datetime.hour) +
#                 'h' + '{:02d}'.format(Time(end_time_target, out_subfmt='date_hm').datetime.minute) +
#                 ' : ' + str(table_schedule['target'][j]))
#             run.bold = True
#             font = run.font
#             font.size = Pt(12)
#             font.color.rgb = RGBColor(0, 0, 0)
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(0)
#             par_format.space_after = Pt(0)
#             run = par.add_run('  Note: Prio_target                                         ')
#             font = run.font
#             font.size = Pt(10)
#             font.color.rgb = RGBColor(0, 0, 0)
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(0)
#             par_format.space_after = Pt(0)
#             run = par.add_run(
#                 '  SPECULOOS : ' + str(df['nb_hours_surved'][idx_target].data.data[0]*u.hour) + ' of obs over ' + str(
#                     df['nb_hours_threshold'][idx_target].data.data[0]*u.hour))
#             font = run.font
#             font.size = Pt(10)
#             font.color.rgb = RGBColor(0, 0, 0)
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(0)
#             par_format.space_after = Pt(0)
#             run = par.add_run('Jmag= ' + str(df['J'][idx_target].data.data[0]) + ',  SpT= ' + str(
#                 df['SpT'][idx_target].data[0]))  # + ', Moon at ' + str(dist_moon))
#             font = run.font
#             font.size = Pt(12)
#             font.color.rgb = RGBColor(0, 0, 0)
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(0)
#             par_format.space_after = Pt(3)
#             run = par.add_run(' RA = ' + str('{:02d}'.format(int(coords.ra.hms[0]))) + " " +
#                               str('{:02d}'.format(int(coords.ra.hms[1]))) + " " +
#                               str('{:05.3f}'.format(round(coords.ra.hms[2], 3))) +
#                               ', DEC = ' + str('{:02d}'.format(int(coords.dec.dms[0]))) + " " +
#                               str('{:02d}'.format(int(abs(coords.dec.dms[1])))) +
#                               " " + str('{:05.3f}'.format(round(abs(coords.dec.dms[2]), 3))) +
#                               ', ' + str(config[2:-2]).replace('\'', ' '))
#             font = run.font
#             font.size = Pt(12)
#             font.color.rgb = RGBColor(0, 0, 0)
#             par = doc.add_paragraph()
#             par_format = par.paragraph_format
#             par_format.space_beFore = Pt(16)
#             par_format.space_after = Pt(0)
#
#     font = run.font
#     font.size = Pt(12)
#     font.color.rgb = RGBColor(0, 0, 0)
#     if telescope == 'TS_La_Silla':
#         doc.save(path_spock + '/TRAPPIST_schedules_docx/TS_' +
#                  Time(date_range[0], out_subfmt='date').value.replace('-', '') + '_to_' +
#                  Time(date_range[1], out_subfmt='date').value .replace('-', '') + '.docx')
#     if telescope == 'TN_Oukaimeden':
#         doc.save(path_spock + '/TRAPPIST_schedules_docx/TN_' +
#                  Time(date_range[0], out_subfmt='date').value.replace('-', '') + '_to_' +
#                  Time(date_range[1], out_subfmt='date').value .replace('-', '') + '.docx')
